"""
Document Filler - Excel/Wordフォーマットにドラフト内容を入力する。

Cloud Run（Linux）対応。openpyxlとpython-docxを使用。
"""

import logging
import os
import shutil
from typing import Any, Dict, List, Optional, Tuple
from datetime import datetime


class DocumentFiller:
    """
    Excel/Wordフォーマットにドラフト内容を入力する。
    Cloud Run（Linux）対応。
    """
    
    def __init__(self, output_dir: str = None):
        """
        Args:
            output_dir: 出力ディレクトリ（デフォルト: /tmp/filled_documents）
        """
        self.output_dir = output_dir or "/tmp/filled_documents"
        self.logger = logging.getLogger(__name__)
        
        # 出力ディレクトリを作成
        os.makedirs(self.output_dir, exist_ok=True)
    
    def _get_existing_font_style(self, paragraph):
        """
        段落から既存のフォントスタイルを取得する。
        最初のrunのスタイルを使用。
        
        Returns:
            Dict with font_name, font_size, bold, italic
        """
        style = {
            "font_name": None,
            "font_size": None,
            "bold": None,
            "italic": None
        }
        
        try:
            if paragraph.runs:
                first_run = paragraph.runs[0]
                if first_run.font:
                    style["font_name"] = first_run.font.name
                    style["font_size"] = first_run.font.size
                    style["bold"] = first_run.font.bold
                    style["italic"] = first_run.font.italic
            
            # runがない場合やフォントが取れない場合、段落スタイルから取得を試みる
            if style["font_name"] is None and paragraph.style and paragraph.style.font:
                style["font_name"] = paragraph.style.font.name
                style["font_size"] = paragraph.style.font.size
                style["bold"] = paragraph.style.font.bold
                style["italic"] = paragraph.style.font.italic
        except Exception as e:
            self.logger.debug(f"[DOC_FILLER] Could not get font style: {e}")
        
        return style
    
    def _add_run_with_style(self, paragraph, text: str, style: Dict = None):
        """
        指定されたスタイルでrunを追加する。
        
        Args:
            paragraph: 対象の段落
            text: 追加するテキスト
            style: フォントスタイル辞書（_get_existing_font_styleの戻り値）
        """
        run = paragraph.add_run(text)
        
        if style:
            try:
                if style.get("font_name"):
                    run.font.name = style["font_name"]
                if style.get("font_size"):
                    run.font.size = style["font_size"]
                if style.get("bold") is not None:
                    run.font.bold = style["bold"]
                if style.get("italic") is not None:
                    run.font.italic = style["italic"]
            except Exception as e:
                self.logger.debug(f"[DOC_FILLER] Could not apply font style: {e}")
        
        return run
    
    def _clear_and_add_with_style(self, paragraph, text: str):
        """
        段落をクリアしてスタイルを保持したままテキストを追加する。
        """
        # 既存のスタイルを保存
        style = self._get_existing_font_style(paragraph)
        
        # 段落をクリア
        paragraph.clear()
        
        # スタイルを適用してテキストを追加
        return self._add_run_with_style(paragraph, text, style)
    
    def fill_document(
        self, 
        file_path: str, 
        field_values: Dict[str, str],
        user_id: str = None
    ) -> Tuple[Optional[str], str]:
        """
        ファイル形式を自動判定してドキュメントに入力する。
        
        Args:
            file_path: テンプレートファイルのパス
            field_values: {field_id: 入力値} のマッピング
            user_id: ユーザーID（出力ファイル名に使用）
            
        Returns:
            (出力ファイルパス, メッセージ)
            入力失敗時は (None, エラーメッセージ)
        """
        if not os.path.exists(file_path):
            return None, f"ファイルが見つかりません: {file_path}"
        
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext in ['.xlsx', '.xlsm', '.xls']:
                return self.fill_excel(file_path, field_values, user_id)
            elif ext in ['.docx', '.doc']:
                return self.fill_word(file_path, field_values, user_id)
            else:
                return None, f"未対応のファイル形式: {ext}"
        except Exception as e:
            self.logger.error(f"[DOC_FILLER] Fill failed: {e}")
            return None, f"入力エラー: {e}"
    
    def fill_excel(
        self, 
        file_path: str, 
        field_values: Dict[str, str],
        user_id: str = None
    ) -> Tuple[Optional[str], str]:
        """
        openpyxlでExcelに入力する。
        
        Args:
            file_path: テンプレートExcelファイルのパス
            field_values: {field_id: 入力値}
                field_id形式: "シート名_行_列" (例: "Sheet1_5_3")
            user_id: ユーザーID
            
        Returns:
            (出力ファイルパス, メッセージ)
        """
        try:
            import openpyxl
            from openpyxl.comments import Comment
        except ImportError:
            return None, "openpyxlがインストールされていません"
        
        try:
            # テンプレートをコピー
            output_path = self._create_output_path(file_path, user_id, "xlsx")
            shutil.copy2(file_path, output_path)
            
            # ファイルを開いて編集
            wb = openpyxl.load_workbook(output_path)
            filled_count = 0
            concern_count = 0
            
            for field_id, field_data in field_values.items():
                # 新形式と旧形式の両方に対応
                if isinstance(field_data, dict):
                    value = field_data.get("value", "")
                    concern_type = field_data.get("concern_type", "none")
                    concern_reason = field_data.get("concern_reason", "")
                    field_name = field_data.get("field_name", field_id)
                else:
                    value = field_data
                    concern_type = "none"
                    concern_reason = ""
                    field_name = field_id
                
                if not value:
                    continue
                
                try:
                    # field_id: "シート名_行_列"
                    parts = field_id.rsplit('_', 2)
                    if len(parts) != 3:
                        self.logger.warning(f"[DOC_FILLER] Invalid field_id format: {field_id}")
                        continue
                    
                    sheet_name, row_str, col_str = parts
                    row = int(row_str)
                    col = int(col_str)
                    
                    if sheet_name not in wb.sheetnames:
                        self.logger.warning(f"[DOC_FILLER] Sheet not found: {sheet_name}")
                        continue
                    
                    sheet = wb[sheet_name]
                    cell = sheet.cell(row=row, column=col, value=value)
                    filled_count += 1
                    
                    # 懸念点がある場合はコメントを追加
                    if concern_type != "none" and concern_reason:
                        comment_text = self._get_concern_comment_text(concern_type, concern_reason, field_name)
                        cell.comment = Comment(comment_text, "Shadow Director AI")
                        concern_count += 1
                        self.logger.debug(f"[DOC_FILLER] Added comment to {field_id}: {concern_type}")
                    
                except (ValueError, IndexError) as e:
                    self.logger.warning(f"[DOC_FILLER] Error filling field {field_id}: {e}")
            
            wb.save(output_path)
            wb.close()
            
            self.logger.info(f"[DOC_FILLER] Filled {filled_count} fields in Excel, {concern_count} comments added")
            
            if filled_count == 0:
                return None, "入力できるフィールドがありませんでした"
            
            message = f"Excelに{filled_count}項目を入力しました"
            if concern_count > 0:
                message += f"（{concern_count}件の懸念点コメント付き）"
            
            return output_path, message
            
        except Exception as e:
            self.logger.error(f"[DOC_FILLER] Excel fill error: {e}")
            return None, f"Excel入力エラー: {e}"
    
    def fill_word(
        self, 
        file_path: str, 
        field_values: Dict[str, Any],
        user_id: str = None
    ) -> Tuple[Optional[str], str]:
        """
        python-docxでWordに入力する。
        
        Args:
            file_path: テンプレートWordファイルのパス
            field_values: フィールド値のマッピング
                新形式: {field_id: {"value": str, "input_pattern": str, "location": dict}}
                旧形式: {field_id: str} （互換性維持）
            user_id: ユーザーID
            
        Returns:
            (出力ファイルパス, メッセージ)
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
        except ImportError:
            return None, "python-docxがインストールされていません"
        
        try:
            # テンプレートをコピー
            output_path = self._create_output_path(file_path, user_id, "docx")
            shutil.copy2(file_path, output_path)
            
            # ファイルを開いて編集
            doc = Document(output_path)
            filled_count = 0
            concern_count = 0
            
            # コメント用のパーツを初期化
            self._init_comments_part(doc)
            
            # 懸念点があるフィールドの情報を蓄積（コメント追加用）
            concerns_to_add = []
            
            for field_id, field_data in field_values.items():
                # 新形式と旧形式の両方に対応
                if isinstance(field_data, dict):
                    value = field_data.get("value", "")
                    input_pattern = field_data.get("input_pattern", "inline")
                    location = field_data.get("location", {})
                    input_length_type = field_data.get("input_length_type", "unknown")
                    concern_type = field_data.get("concern_type", "none")
                    concern_reason = field_data.get("concern_reason", "")
                    field_name = field_data.get("field_name", field_id)
                else:
                    # 旧形式（文字列のみ）
                    value = field_data
                    input_pattern = "inline"
                    location = {}
                    input_length_type = "unknown"
                    concern_type = "none"
                    concern_reason = ""
                    field_name = field_id
                
                if not value:
                    continue
                
                # 懸念点がある場合は後でコメントを追加
                has_concern = concern_type != "none" and concern_reason
                
                try:
                    filled = False
                    target_paragraph = None
                    
                    if field_id.startswith("table"):
                        # テーブルセル: "tableN_行_列" - input_length_typeを考慮
                        filled, target_paragraph = self._fill_word_table_cell_with_para(doc, field_id, value, input_length_type)
                    elif field_id.startswith("para_"):
                        # 段落: "para_N" - 入力パターン情報を使用
                        filled, target_paragraph = self._fill_word_paragraph_with_pattern_and_para(doc, field_id, value, input_pattern, location)
                    else:
                        self.logger.warning(f"[DOC_FILLER] Unknown field_id format: {field_id}")
                    
                    if filled:
                        filled_count += 1
                        self.logger.debug(f"[DOC_FILLER] Filled {field_id} with pattern '{input_pattern}'")
                        
                        # 懸念点がある場合、コメント追加対象としてリストに追加
                        if has_concern and target_paragraph is not None:
                            concern_count += 1
                            concerns_to_add.append({
                                "paragraph": target_paragraph,
                                "field_name": field_name,
                                "concern_type": concern_type,
                                "concern_reason": concern_reason
                            })
                        
                except Exception as e:
                    self.logger.warning(f"[DOC_FILLER] Error filling field {field_id}: {e}")
            
            # 懸念点コメントを追加
            for idx, concern in enumerate(concerns_to_add):
                try:
                    comment_text = self._get_concern_comment_text(
                        concern["concern_type"], 
                        concern["concern_reason"], 
                        concern["field_name"]
                    )
                    self._add_word_native_comment(
                        doc, 
                        concern["paragraph"], 
                        comment_text, 
                        idx
                    )
                except Exception as e:
                    self.logger.warning(f"[DOC_FILLER] Failed to add comment: {e}")
            
            doc.save(output_path)
            
            # コメントをZIPファイルに注入（python-docxの制限を回避）
            if concerns_to_add and hasattr(doc, '_comments_element'):
                self._inject_comments_to_docx(output_path, doc._comments_element)
            
            self.logger.info(f"[DOC_FILLER] Filled {filled_count} fields in Word, {concern_count} comments added")
            
            if filled_count == 0:
                return None, "入力できるフィールドがありませんでした"
            
            message = f"Wordに{filled_count}項目を入力しました"
            if concern_count > 0:
                message += f"（{concern_count}件のコメント付き）"
            
            return output_path, message
            
        except Exception as e:
            self.logger.error(f"[DOC_FILLER] Word fill error: {e}")
            return None, f"Word入力エラー: {e}"

    
    def _fill_word_table_cell(self, doc, field_id: str, value: str, input_length_type: str = "unknown") -> bool:
        """
        Wordテーブルセルに入力。
        
        Args:
            doc: Wordドキュメント
            field_id: フィールドID（"tableN_行_列"形式）
            value: 入力値
            input_length_type: "short"（短文）, "long"（長文）, "unknown"
        """
        try:
            # "tableN_行_列" をパース
            parts = field_id.split('_')
            if len(parts) != 3:
                return False
            
            table_part = parts[0]  # "tableN"
            row = int(parts[1])
            col = int(parts[2])
            table_idx = int(table_part.replace("table", ""))
            
            if table_idx >= len(doc.tables):
                self.logger.warning(f"[DOC_FILLER] Table {table_idx} not found")
                return False
            
            table = doc.tables[table_idx]
            
            if row >= len(table.rows):
                self.logger.warning(f"[DOC_FILLER] Row {row} not found in table {table_idx}")
                return False
            
            cells = table.rows[row].cells
            if col >= len(cells):
                self.logger.warning(f"[DOC_FILLER] Col {col} not found in table {table_idx}, row {row}")
                return False
            
            cell = cells[col]
            
            # 長文の場合、テーブルセル内に収まるように処理
            # 短い場合はそのまま、長い場合は文字数を制限して...を付ける
            if input_length_type == "short" and len(value) > 50:
                # 短文フィールドに長いテキストが来た場合、切り詰める
                value = value[:47] + "..."
                self.logger.debug(f"[DOC_FILLER] Trimmed long value for short field: {field_id}")
            
            # 既存テキストをクリアして新しいテキストを設定
            # フォントスタイルを保持する
            if cell.paragraphs:
                para = cell.paragraphs[0]
                self._clear_and_add_with_style(para, value)
            else:
                cell.text = value
            
            return True
            
        except Exception as e:
            self.logger.warning(f"[DOC_FILLER] Table cell fill error: {e}")
            return False
    
    def _fill_word_paragraph(self, doc, field_id: str, value: str) -> bool:
        """
        Word段落に入力（プレースホルダーを置換）。
        
        対応する入力タイプ:
        - inline: コロン後に入力を追加
        - next_line: 段落全体を入力値で置換
        - underline: 下線プレースホルダーを置換
        - bracket: 空括弧プレースホルダーを置換
        """
        try:
            # "para_N" をパース
            para_idx = int(field_id.replace("para_", ""))
            
            if para_idx >= len(doc.paragraphs):
                self.logger.warning(f"[DOC_FILLER] Paragraph {para_idx} not found")
                return False
            
            para = doc.paragraphs[para_idx]
            original_text = para.text
            
            import re
            
            # パターン1: 下線プレースホルダー「____」「＿＿＿」を置換
            new_text = re.sub(r'[_＿]{3,}', value, original_text)
            if new_text != original_text:
                para.clear()
                para.add_run(new_text)
                return True
            
            # パターン2: 空括弧プレースホルダー「（　）」を置換
            new_text = re.sub(r'[(（]\s*[　\s]*[)）]', f'（{value}）', original_text)
            if new_text != original_text:
                para.clear()
                para.add_run(new_text)
                return True
            
            # パターン3: 括弧付きヒント「（入力してください）」を置換
            new_text = re.sub(r'[(（][^)）]+[)）]', f'（{value}）', original_text)
            if new_text != original_text:
                para.clear()
                para.add_run(new_text)
                return True
            
            # パターン4: コロン終端の場合、コロン後に入力を追加
            colon_match = re.match(r'^(.+?[:：])\s*$', original_text)
            if colon_match:
                new_text = f"{colon_match.group(1)} {value}"
                para.clear()
                para.add_run(new_text)
                return True
            
            # パターン5: コロンがある場合、コロン後を置換
            colon_replace_match = re.match(r'^(.+?[:：])\s*(.*)$', original_text)
            if colon_replace_match:
                prefix = colon_replace_match.group(1)
                current_value = colon_replace_match.group(2).strip()
                
                # 現在の値が空、空白のみ、またはヒント（括弧付き）の場合に置換
                if not current_value or re.match(r'^[　\s]+$', current_value) or re.match(r'^[（(].+[)）]$', current_value):
                    new_text = f"{prefix} {value}"
                    para.clear()
                    para.add_run(new_text)
                    return True
            
            # パターン6: 次行入力の場合（段落が比較的空の場合）、テキスト全体を置換
            if len(original_text.strip()) < 10:
                para.clear()
                para.add_run(value)
                return True
            
            # 上記いずれにも該当しない場合、段落末尾に追加
            para.add_run(f"\n{value}")
            return True
            
        except Exception as e:
            self.logger.warning(f"[DOC_FILLER] Paragraph fill error: {e}")
            return False
    
    def _fill_word_paragraph_with_pattern(
        self, 
        doc, 
        field_id: str, 
        value: str, 
        input_pattern: str,
        location: Dict[str, Any]
    ) -> bool:
        """
        VLMで検出された入力パターンに基づいてWord段落に入力する。
        
        Args:
            doc: Wordドキュメント
            field_id: フィールドID（"para_N"形式）
            value: 入力値
            input_pattern: 入力パターン（"inline", "next_line", "underline", "bracket"）
            location: 位置情報（paragraph_idx, label_paragraph_idx等）
            
        Returns:
            入力成功かどうか
        """
        try:
            import re
            
            # パターンに応じた処理
            para_idx = location.get("paragraph_idx")
            if para_idx is None:
                # field_idからパース
                para_idx = int(field_id.replace("para_", ""))
            
            if para_idx >= len(doc.paragraphs):
                self.logger.warning(f"[DOC_FILLER] Paragraph {para_idx} not found")
                return False
            
            para = doc.paragraphs[para_idx]
            original_text = para.text
            
            # パターン別の入力処理
            # 事前にスタイルを取得
            style = self._get_existing_font_style(para)
            
            if input_pattern == "next_line":
                # 次行入力: 段落全体を入力値で置換（スタイル保持）
                # この段落がラベルの次の段落なので、内容を完全に置き換える
                para.clear()
                self._add_run_with_style(para, value, style)
                self.logger.debug(f"[DOC_FILLER] Applied next_line pattern to para {para_idx}")
                return True
            
            elif input_pattern == "underline":
                # 下線プレースホルダー「____」を置換（スタイル保持）
                new_text = re.sub(r'[_＿]{2,}', value, original_text)
                if new_text != original_text:
                    para.clear()
                    self._add_run_with_style(para, new_text, style)
                    self.logger.debug(f"[DOC_FILLER] Applied underline pattern to para {para_idx}")
                    return True
                # 下線がない場合はinlineとして処理
                input_pattern = "inline"
            
            elif input_pattern == "bracket":
                # 括弧プレースホルダー「（　）」「（入力してください）」を置換（スタイル保持）
                # まず空括弧を試す
                new_text = re.sub(r'[(（]\s*[　\s]*[)）]', f'（{value}）', original_text)
                if new_text != original_text:
                    para.clear()
                    self._add_run_with_style(para, new_text, style)
                    self.logger.debug(f"[DOC_FILLER] Applied bracket pattern (empty) to para {para_idx}")
                    return True
                # ヒント付き括弧を試す
                new_text = re.sub(r'[(（][^)）]+[)）]', f'（{value}）', original_text)
                if new_text != original_text:
                    para.clear()
                    self._add_run_with_style(para, new_text, style)
                    self.logger.debug(f"[DOC_FILLER] Applied bracket pattern (with hint) to para {para_idx}")
                    return True
                # 括弧がない場合はinlineとして処理
                input_pattern = "inline"
            
            # inlineパターン（デフォルト）
            if input_pattern == "inline":
                # コロン後に入力を追加/置換（スタイル保持）
                colon_match = re.match(r'^(.+?[:：])\s*(.*)$', original_text)
                if colon_match:
                    prefix = colon_match.group(1)
                    current_value = colon_match.group(2).strip()
                    
                    # 現在の値が空、空白のみ、下線、またはヒント（括弧付き）の場合に置換
                    if (not current_value or 
                        re.match(r'^[　\s]+$', current_value) or 
                        re.match(r'^[_＿]+$', current_value) or
                        re.match(r'^[（(].+[)）]$', current_value)):
                        new_text = f"{prefix} {value}"
                        para.clear()
                        self._add_run_with_style(para, new_text, style)
                        self.logger.debug(f"[DOC_FILLER] Applied inline pattern to para {para_idx}")
                        return True
                    else:
                        # 既存の値がある場合は置き換える
                        new_text = f"{prefix} {value}"
                        para.clear()
                        self._add_run_with_style(para, new_text, style)
                        self.logger.debug(f"[DOC_FILLER] Replaced existing value with inline pattern in para {para_idx}")
                        return True
                
                # コロンがない場合は段落末尾に追加（スタイル保持）
                self._add_run_with_style(para, f" {value}", style)
                self.logger.debug(f"[DOC_FILLER] Appended value to para {para_idx} (no colon found)")
                return True
            
            # 不明なパターンの場合はフォールバックとして既存メソッドを使用
            self.logger.warning(f"[DOC_FILLER] Unknown pattern '{input_pattern}', using fallback")
            return self._fill_word_paragraph(doc, field_id, value)
            
        except Exception as e:
            self.logger.warning(f"[DOC_FILLER] Pattern-based paragraph fill error: {e}")
            # フォールバックとして既存メソッドを試す
            try:
                return self._fill_word_paragraph(doc, field_id, value)
            except:
                return False
    
    def _create_output_path(self, original_path: str, user_id: str, ext: str) -> str:
        """出力ファイルパスを生成"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        original_name = os.path.splitext(os.path.basename(original_path))[0]
        
        user_dir = os.path.join(self.output_dir, user_id or "default")
        os.makedirs(user_dir, exist_ok=True)
        
        output_name = f"{original_name}_filled_{timestamp}.{ext}"
        return os.path.join(user_dir, output_name)
    
    def cleanup_old_files(self, max_age_hours: int = 24):
        """古い出力ファイルを削除"""
        try:
            import time
            
            current_time = time.time()
            max_age_seconds = max_age_hours * 3600
            
            for root, dirs, files in os.walk(self.output_dir):
                for filename in files:
                    file_path = os.path.join(root, filename)
                    file_age = current_time - os.path.getmtime(file_path)
                    
                    if file_age > max_age_seconds:
                        os.remove(file_path)
                        self.logger.info(f"[DOC_FILLER] Cleaned up old file: {filename}")
                        
        except Exception as e:
            self.logger.warning(f"[DOC_FILLER] Cleanup error: {e}")
    
    def _get_concern_comment_text(self, concern_type: str, concern_reason: str, field_name: str) -> str:
        """
        懸念点タイプに応じたコメントテキストを生成する。
        
        Args:
            concern_type: 懸念点タイプ（missing_info, uncertain, length_exceeded, truncated）
            concern_reason: 懸念点の理由
            field_name: フィールド名
            
        Returns:
            コメントテキスト
        """
        type_labels = {
            "missing_info": "⚠️ 情報不足",
            "uncertain": "❓ 要確認",
            "length_exceeded": "📏 文字数超過",
            "truncated": "✂️ 回答省略"
        }
        
        type_label = type_labels.get(concern_type, "⚠️ 懸念あり")
        
        comment = f"""【{type_label}】
項目: {field_name}
理由: {concern_reason}

※ 内容をご確認のうえ、必要に応じて修正してください。
(自動生成: Shadow Director AI)"""
        
        return comment
    
    def _add_word_concerns_section(self, doc, concerns_list: list):
        """
        Wordドキュメント末尾に懸念点一覧セクションを追加する。
        
        Args:
            doc: Wordドキュメント
            concerns_list: 懸念点情報のリスト
        """
        try:
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            self.logger.warning("[DOC_FILLER] Failed to import docx components for concerns section")
            return
        
        try:
            # 区切り線として空白行を追加
            doc.add_paragraph("")
            doc.add_paragraph("─" * 40)
            
            # タイトル段落
            title_para = doc.add_paragraph()
            title_run = title_para.add_run("📋 Shadow Director AI - 懸念点一覧")
            title_run.bold = True
            title_run.font.size = Pt(12)
            
            # 説明
            desc_para = doc.add_paragraph()
            desc_run = desc_para.add_run("以下の項目については、内容をご確認のうえ必要に応じて修正してください。")
            desc_run.font.size = Pt(10)
            desc_run.font.color.rgb = RGBColor(100, 100, 100)
            
            # 懸念点リスト
            type_labels = {
                "missing_info": "⚠️ 情報不足",
                "uncertain": "❓ 要確認",
                "length_exceeded": "📏 文字数超過",
                "truncated": "✂️ 回答省略"
            }
            
            for concern in concerns_list:
                number = concern["number"]
                field_name = concern["field_name"]
                concern_type = concern["concern_type"]
                concern_reason = concern["concern_reason"]
                
                type_label = type_labels.get(concern_type, "⚠️ 懸念あり")
                
                item_para = doc.add_paragraph()
                item_run = item_para.add_run(f"[※{number}] 【{type_label}】{field_name}")
                item_run.bold = True
                item_run.font.size = Pt(10)
                
                reason_para = doc.add_paragraph()
                reason_run = reason_para.add_run(f"    → {concern_reason}")
                reason_run.font.size = Pt(9)
                reason_run.font.color.rgb = RGBColor(80, 80, 80)
            
            # フッター
            doc.add_paragraph("")
            footer_para = doc.add_paragraph()
            footer_run = footer_para.add_run("※ この懸念点一覧は提出前に削除してください。")
            footer_run.font.size = Pt(8)
            footer_run.font.color.rgb = RGBColor(150, 150, 150)
            footer_run.italic = True
            
            self.logger.info(f"[DOC_FILLER] Added concerns section with {len(concerns_list)} items")
            
        except Exception as e:
            self.logger.warning(f"[DOC_FILLER] Failed to add concerns section: {e}")
    
    def _init_comments_part(self, doc):
        """
        ドキュメントにコメントパーツを初期化する。
        python-docxは標準でcomments.xmlを作成しないため、OOXMLで追加する。
        """
        try:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            from docx.parts.document import DocumentPart
            
            # コメントパーツが既に存在するかチェック
            document_part = doc.part
            
            # comments関係を探す
            for rel in document_part.rels.values():
                if 'comments' in rel.reltype:
                    self.logger.debug("[DOC_FILLER] Comments part already exists")
                    return
            
            self.logger.debug("[DOC_FILLER] Comments part initialized (will be created on save if needed)")
            
        except Exception as e:
            self.logger.debug(f"[DOC_FILLER] Comments part init skipped: {e}")
    
    def _add_word_native_comment(self, doc, paragraph, comment_text: str, comment_id: int):
        """
        Wordドキュメントにネイティブコメントを追加する（OOXML直接操作）。
        
        Args:
            doc: Wordドキュメント
            paragraph: コメントを追加する段落
            comment_text: コメントテキスト
            comment_id: コメントID（0から始まる連番）
        """
        try:
            from docx.oxml.ns import qn, nsmap
            from docx.oxml import OxmlElement
            from datetime import datetime
            
            # Word namespace
            w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            
            # コメントIDを文字列に
            cid = str(comment_id)
            
            # 段落にコメント参照マーカーを追加
            # commentRangeStart要素を作成
            comment_range_start = OxmlElement('w:commentRangeStart')
            comment_range_start.set(qn('w:id'), cid)
            
            # commentRangeEnd要素を作成
            comment_range_end = OxmlElement('w:commentRangeEnd')
            comment_range_end.set(qn('w:id'), cid)
            
            # commentReference要素を作成（w:r 内に入れ、rPrも必須）
            comment_ref_run = OxmlElement('w:r')
            
            # ランプロパティ (w:rPr) を追加 - フォントサイズのみ設定（スタイル参照は避ける）
            run_props = OxmlElement('w:rPr')
            # コメント参照は通常8pt程度の上付き文字
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '16')  # 8pt = 16 half-points
            run_props.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '16')
            run_props.append(szCs)
            comment_ref_run.append(run_props)
            
            # commentReference を追加
            comment_ref = OxmlElement('w:commentReference')
            comment_ref.set(qn('w:id'), cid)
            comment_ref_run.append(comment_ref)
            
            # 段落の最初と最後にマーカーを挿入
            para_element = paragraph._p
            
            # pPr（段落プロパティ）がある場合、その後に挿入
            # pPrがない場合は最初に挿入
            pPr = para_element.find(qn('w:pPr'))
            if pPr is not None:
                # pPrの次に挿入
                pPr_index = list(para_element).index(pPr)
                para_element.insert(pPr_index + 1, comment_range_start)
            elif len(para_element) > 0:
                para_element.insert(0, comment_range_start)
            else:
                para_element.append(comment_range_start)
            
            # 段落の最後にcommentRangeEndとcommentReferenceを追加
            para_element.append(comment_range_end)
            para_element.append(comment_ref_run)
            
            # comments.xmlにコメント本体を追加
            self._add_comment_to_comments_part(doc, cid, comment_text)
            
            self.logger.debug(f"[DOC_FILLER] Added native comment {cid} to paragraph")
            
        except Exception as e:
            self.logger.warning(f"[DOC_FILLER] Failed to add native comment: {e}")
            # フォールバック: 段落末尾にコメントテキストを追加
            try:
                from docx.shared import Pt, RGBColor
                run = paragraph.add_run(f" [※コメント: {comment_text[:50]}...]")
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(128, 128, 128)
                run.italic = True
            except:
                pass
    
    def _add_comment_to_comments_part(self, doc, comment_id: str, comment_text: str):
        """
        comments.xmlにコメント本体を追加する。
        
        Args:
            doc: Wordドキュメント
            comment_id: コメントID
            comment_text: コメントテキスト
        """
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from datetime import datetime
            from lxml import etree
            
            document_part = doc.part
            
            # comments要素を取得または作成
            comments_element = self._get_or_create_comments_element(doc)
            if comments_element is None:
                self.logger.warning("[DOC_FILLER] Could not get/create comments element")
                return
            
            # コメント要素を作成
            comment = OxmlElement('w:comment')
            comment.set(qn('w:id'), comment_id)
            comment.set(qn('w:author'), 'Shadow Director AI')
            comment.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
            comment.set(qn('w:initials'), 'SD')
            
            # コメント本文を段落として追加
            # 複数行がある場合は分割
            lines = comment_text.split('\n')
            
            for line in lines:
                comment_para = OxmlElement('w:p')
                
                # 段落プロパティ (w:pPr) を追加
                para_props = OxmlElement('w:pPr')
                comment_para.append(para_props)
                
                if line.strip():
                    comment_run = OxmlElement('w:r')
                    
                    # ランプロパティ (w:rPr) を追加 - これが欠落するとエラーになる
                    run_props = OxmlElement('w:rPr')
                    run_props_lang = OxmlElement('w:lang')
                    run_props_lang.set(qn('w:val'), 'ja-JP')
                    run_props.append(run_props_lang)
                    comment_run.append(run_props)
                    
                    # テキスト要素を作成
                    comment_text_elem = OxmlElement('w:t')
                    comment_text_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    comment_text_elem.text = line
                    comment_run.append(comment_text_elem)
                    
                    comment_para.append(comment_run)
                
                comment.append(comment_para)
            
            # commentsに追加
            comments_element.append(comment)
            
            self.logger.debug(f"[DOC_FILLER] Added comment {comment_id} to comments.xml")
            
        except Exception as e:
            self.logger.warning(f"[DOC_FILLER] Failed to add comment to comments part: {e}")
    
    def _get_or_create_comments_element(self, doc):
        """
        comments.xmlのルート要素を取得または作成する。
        """
        try:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from lxml import etree
            import zipfile
            import io
            
            document_part = doc.part
            
            # 既存のコメントパーツを探す
            for rel in document_part.rels.values():
                if 'comments' in rel.reltype:
                    return rel.target_part._element
            
            # コメントパーツがない場合は、ドキュメント自体にcommentsを埋め込む方式を試す
            # (python-docxの制限により、新規パーツ追加は複雑なため)
            
            # 代替手段: 属性としてcommentsを保持
            if not hasattr(doc, '_comments_element'):
                # 新しいcomments要素を作成
                comments = OxmlElement('w:comments')
                doc._comments_element = comments
            
            return doc._comments_element
            
        except Exception as e:
            self.logger.warning(f"[DOC_FILLER] Failed to get/create comments element: {e}")
            return None
    
    def _fill_word_table_cell_with_para(self, doc, field_id: str, value: str, input_length_type: str = "unknown") -> Tuple[bool, Optional[Any]]:
        """
        Wordテーブルセルに入力し、対象の段落を返す。
        
        Args:
            doc: Wordドキュメント
            field_id: フィールドID（"tableN_行_列"形式）
            value: 入力値
            input_length_type: "short"（短文）, "long"（長文）, "unknown"
            
        Returns:
            (成功フラグ, 対象段落)
        """
        try:
            # "tableN_行_列" をパース
            parts = field_id.split('_')
            if len(parts) != 3:
                return False, None
            
            table_part = parts[0]  # "tableN"
            row = int(parts[1])
            col = int(parts[2])
            table_idx = int(table_part.replace("table", ""))
            
            if table_idx >= len(doc.tables):
                self.logger.warning(f"[DOC_FILLER] Table {table_idx} not found")
                return False, None
            
            table = doc.tables[table_idx]
            
            if row >= len(table.rows):
                self.logger.warning(f"[DOC_FILLER] Row {row} not found in table {table_idx}")
                return False, None
            
            cells = table.rows[row].cells
            if col >= len(cells):
                self.logger.warning(f"[DOC_FILLER] Col {col} not found in table {table_idx}, row {row}")
                return False, None
            
            cell = cells[col]
            
            # 長文の場合、テーブルセル内に収まるように処理
            if input_length_type == "short" and len(value) > 50:
                value = value[:47] + "..."
                self.logger.debug(f"[DOC_FILLER] Trimmed long value for short field: {field_id}")
            
            # 既存テキストをクリアして新しいテキストを設定
            target_para = None
            if cell.paragraphs:
                para = cell.paragraphs[0]
                self._clear_and_add_with_style(para, value)
                target_para = para
            else:
                cell.text = value
                if cell.paragraphs:
                    target_para = cell.paragraphs[0]
            
            return True, target_para
            
        except Exception as e:
            self.logger.warning(f"[DOC_FILLER] Table cell fill error: {e}")
            return False, None
    
    def _fill_word_paragraph_with_pattern_and_para(
        self, 
        doc, 
        field_id: str, 
        value: str, 
        input_pattern: str,
        location: Dict[str, Any]
    ) -> Tuple[bool, Optional[Any]]:
        """
        VLMで検出された入力パターンに基づいてWord段落に入力し、段落を返す。
        
        Args:
            doc: Wordドキュメント
            field_id: フィールドID（"para_N"形式）
            value: 入力値
            input_pattern: 入力パターン
            location: 位置情報
            
        Returns:
            (成功フラグ, 対象段落)
        """
        try:
            import re
            
            # パターンに応じた処理
            para_idx = location.get("paragraph_idx")
            if para_idx is None:
                para_idx = int(field_id.replace("para_", ""))
            
            if para_idx >= len(doc.paragraphs):
                self.logger.warning(f"[DOC_FILLER] Paragraph {para_idx} not found")
                return False, None
            
            para = doc.paragraphs[para_idx]
            original_text = para.text
            
            # スタイルを取得
            style = self._get_existing_font_style(para)
            
            # 入力処理（既存の_fill_word_paragraph_with_patternと同様）
            filled = False
            
            if input_pattern == "next_line":
                para.clear()
                self._add_run_with_style(para, value, style)
                filled = True
            elif input_pattern == "underline":
                new_text = re.sub(r'[_＿]{2,}', value, original_text)
                if new_text != original_text:
                    para.clear()
                    self._add_run_with_style(para, new_text, style)
                    filled = True
                else:
                    input_pattern = "inline"
            elif input_pattern == "bracket":
                new_text = re.sub(r'[(（]\s*[　\s]*[)）]', f'（{value}）', original_text)
                if new_text != original_text:
                    para.clear()
                    self._add_run_with_style(para, new_text, style)
                    filled = True
                else:
                    new_text = re.sub(r'[(（][^)）]+[)）]', f'（{value}）', original_text)
                    if new_text != original_text:
                        para.clear()
                        self._add_run_with_style(para, new_text, style)
                        filled = True
                    else:
                        input_pattern = "inline"
            
            if input_pattern == "inline":
                colon_match = re.match(r'^(.+?[:：])\s*(.*)$', original_text)
                if colon_match:
                    prefix = colon_match.group(1)
                    new_text = f"{prefix} {value}"
                    para.clear()
                    self._add_run_with_style(para, new_text, style)
                    filled = True
                else:
                    self._add_run_with_style(para, f" {value}", style)
                    filled = True
            
            if not filled:
                # フォールバック
                success = self._fill_word_paragraph(doc, field_id, value)
                return success, para if success else None
            
            return filled, para
            
        except Exception as e:
            self.logger.warning(f"[DOC_FILLER] Pattern-based paragraph fill error: {e}")
            return False, None
    
    def _inject_comments_to_docx(self, docx_path: str, comments_element):
        """
        保存後のdocxファイルにcomments.xmlを注入する。
        
        python-docxはcomments.xmlを保存しないため、ZIPファイル操作で挿入する。
        
        Args:
            docx_path: 保存済みのdocxファイルパス
            comments_element: コメント要素（w:comments）
        """
        try:
            import zipfile
            import tempfile
            from lxml import etree
            
            # コメントがない場合はスキップ
            if len(comments_element) == 0:
                self.logger.debug("[DOC_FILLER] No comments to inject")
                return
            
            # 一時ファイルを作成
            temp_path = docx_path + ".tmp"
            
            # 正しいOOXML形式のcomments.xmlを手動で構築
            comments_xml = self._build_comments_xml(comments_element)
            
            # 既存のdocxを読み込んで新しいファイルに書き出し
            with zipfile.ZipFile(docx_path, 'r') as zin:
                with zipfile.ZipFile(temp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                    for item in zin.infolist():
                        if item.filename == 'word/_rels/document.xml.rels':
                            # リレーションシップファイルにコメント参照を追加
                            content = zin.read(item.filename).decode('utf-8')
                            content = self._add_comments_relationship(content)
                            zout.writestr(item, content.encode('utf-8'))
                        elif item.filename == '[Content_Types].xml':
                            # Content_Typesにコメントタイプを追加
                            content = zin.read(item.filename).decode('utf-8')
                            content = self._add_comments_content_type(content)
                            zout.writestr(item, content.encode('utf-8'))
                        else:
                            # その他のファイルはそのままコピー
                            zout.writestr(item, zin.read(item.filename))
                    
                    # comments.xmlを追加
                    zout.writestr('word/comments.xml', comments_xml.encode('utf-8'))
            
            # 元のファイルを置き換え
            import os
            os.replace(temp_path, docx_path)
            
            self.logger.info(f"[DOC_FILLER] Injected {len(comments_element)} comments to docx")
            
        except Exception as e:
            self.logger.warning(f"[DOC_FILLER] Failed to inject comments: {e}")
            # 一時ファイルがあれば削除
            try:
                if os.path.exists(temp_path):
                    os.remove(temp_path)
            except:
                pass
    
    def _add_comments_relationship(self, rels_content: str) -> str:
        """
        document.xml.relsにコメント参照を追加する。
        シンプルな文字列置換方式で信頼性を向上。
        """
        try:
            # 既にコメント参照がある場合はスキップ
            if 'comments.xml' in rels_content:
                self.logger.debug("[DOC_FILLER] Comments relationship already exists")
                return rels_content
            
            import re
            
            # 既存のrIdを抽出して最大値を取得
            rids = re.findall(r'Id="rId(\d+)"', rels_content)
            max_rid = max([int(r) for r in rids]) if rids else 0
            new_rid = max_rid + 1
            
            # 新しいリレーションシップを構築
            new_rel = f'<Relationship Id="rId{new_rid}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" Target="comments.xml"/>'
            
            # </Relationships>の前に挿入
            if '</Relationships>' in rels_content:
                rels_content = rels_content.replace('</Relationships>', f'{new_rel}</Relationships>')
                self.logger.info(f"[DOC_FILLER] Added comments relationship as rId{new_rid}")
            else:
                self.logger.warning("[DOC_FILLER] Could not find </Relationships> tag")
            
            return rels_content
            
        except Exception as e:
            self.logger.warning(f"[DOC_FILLER] Failed to add comments relationship: {e}")
            return rels_content
    
    def _add_comments_content_type(self, content_types: str) -> str:
        """
        [Content_Types].xmlにコメントのコンテンツタイプを追加する。
        シンプルな文字列置換方式で信頼性を向上。
        """
        try:
            # 既にコメントタイプがある場合はスキップ
            if 'comments.xml' in content_types:
                self.logger.debug("[DOC_FILLER] Comments content type already exists")
                return content_types
            
            # 新しいオーバーライドを構築
            new_override = '<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>'
            
            # </Types>の前に挿入
            if '</Types>' in content_types:
                content_types = content_types.replace('</Types>', f'{new_override}</Types>')
                self.logger.info("[DOC_FILLER] Added comments content type to [Content_Types].xml")
            else:
                self.logger.warning("[DOC_FILLER] Could not find </Types> tag")
            
            return content_types
            
        except Exception as e:
            self.logger.warning(f"[DOC_FILLER] Failed to add comments content type: {e}")
            return content_types
    
    def _build_comments_xml(self, comments_element) -> str:
        """
        コメント要素から正しいOOXML形式のcomments.xmlを構築する。
        
        Args:
            comments_element: コメント要素のリスト
            
        Returns:
            comments.xmlの内容
        """
        from docx.oxml.ns import qn
        import html
        
        xml_parts = [
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
            '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" ',
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" ',
            'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" ',
            'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">',
        ]
        
        for comment in comments_element:
            # コメント属性を取得
            comment_id = comment.get(qn('w:id'), '0')
            author = comment.get(qn('w:author'), 'Shadow Director AI')
            date = comment.get(qn('w:date'), '')
            initials = comment.get(qn('w:initials'), 'SD')
            
            # コメント開始タグ
            xml_parts.append(f'<w:comment w:id="{comment_id}" w:author="{html.escape(author)}" w:date="{date}" w:initials="{initials}">')
            
            # 各段落を処理
            for para in comment:
                if para.tag.endswith('}p') or para.tag == 'w:p':
                    xml_parts.append('<w:p>')
                    xml_parts.append('<w:pPr/>')
                    
                    # ラン要素を処理
                    for run in para:
                        if run.tag.endswith('}r') or run.tag == 'w:r':
                            xml_parts.append('<w:r>')
                            xml_parts.append('<w:rPr><w:lang w:val="ja-JP"/></w:rPr>')
                            
                            # テキスト要素を処理
                            for text_elem in run:
                                if text_elem.tag.endswith('}t') or text_elem.tag == 'w:t':
                                    text_content = text_elem.text or ''
                                    xml_parts.append(f'<w:t xml:space="preserve">{html.escape(text_content)}</w:t>')
                            
                            xml_parts.append('</w:r>')
                    
                    xml_parts.append('</w:p>')
            
            xml_parts.append('</w:comment>')
        
        xml_parts.append('</w:comments>')
        
        return '\n'.join(xml_parts)
