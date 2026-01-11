import os
import logging
import re
from typing import Optional, List, Dict

class FileClassifier:
    """
    ファイルの分類を行うクラス。
    ファイル名やVLM（Vision-Language Model）を使用して、ファイルが助成金申請書、募集要項、あるいは無関係な資料であるかを判定する。
    """
    
    def __init__(self, gemini_client, vlm_model: str = "gemini-3-flash-preview"):
        """
        Args:
            gemini_client: Gemini APIクライアント
            vlm_model: 使用するVLMモデル名
        """
        self.client = gemini_client
        self.vlm_model = vlm_model
    
    def _sanitize_grant_name(self, grant_name: str) -> str:
        """
        grant_nameからユーザーコマンド（「ドラフトを作成して」等）を除去する。
        """
        if not grant_name:
            return ""
        
        # 除去すべきフレーズ（コマンド系）
        remove_phrases = [
            'のドラフトを作成して',
            'ドラフトを作成して',
            'のドラフト作成',
            'ドラフト作成',
            'の申請書を書いて',
            '申請書を書いて',
            'を書いて',
            'について調べて',
            'について詳しく',
            'を調べて',
            'の詳細',
        ]
        
        sanitized = grant_name
        for phrase in remove_phrases:
            sanitized = sanitized.replace(phrase, '')
        
        return sanitized.strip()
        
    def classify_format_file(self, filename: str, file_path: str = None, grant_name: str = None) -> str:
        """
        ファイル名からファイルの用途を判定する。
        VLMによる判定を最優先に行い、できない場合はキーワード判定を行う。
        
        Args:
            filename: ファイル名
            file_path: ファイルパス（VLM解析用）
            grant_name: 助成金名（VLMで関連性を検証するために使用）
            
        Returns:
            ファイルの用途を示す文字列
        """
        fn_lower = filename.lower()
        
        # grant_nameをサニタイズ（ユーザーコマンドを除去）
        if grant_name:
            grant_name = self._sanitize_grant_name(grant_name)
        
        # VLMによる判定を最優先で実施 (User Request: Always use VLM if possible)
        # ファイル内容と助成金名との関連性を厳密にチェックするため
        if file_path and fn_lower.endswith(('.xlsx', '.xls', '.docx', '.doc', '.pdf')):
            vlm_result = self._classify_file_with_vlm(file_path, filename, grant_name)
            if vlm_result:
                return vlm_result
        
        # VLM判定ができなかった場合（非対応ファイル、エラー等）のみ、ファイル名キーワードで簡易判定
        logging.info(f"[DEBUG] VLM判定ができなかったため、ファイル名キーワードで簡易判定")

        # 募集要項・公募要領系（最優先で判定）
        if any(kw in fn_lower for kw in ['募集要項', '公募要領', '応募要項', '公募要項', '募集案内', '公募案内', 'guidelines', 'requirements']):
            return "📋 募集要項（応募条件・審査基準が記載）"
        
        # 交付要綱・規程系
        if any(kw in fn_lower for kw in ['交付要綱', '交付規程', '実施要領', 'ガイドライン', 'guideline', '手引き', '手引']):
            return "📜 交付要綱・ガイドライン（ルール・規程）"
        
        # 記入例系（申請書より先に判定）
        if any(kw in fn_lower for kw in ['記入例', '記載例', '作成例', 'サンプル', 'sample', '見本', '例', 'example']):
            return "📖 記入例・サンプル（参考資料）"
        
        # 申請書・様式系
        if any(kw in fn_lower for kw in ['申請書', '応募書', '様式', 'フォーマット', 'テンプレート', 'template', 'form', '届出', '調書']):
            return "📝 申請書フォーマット（記入が必要）"
        
        # 予算書系
        if any(kw in fn_lower for kw in ['予算', '収支', '経費', 'budget', '見積']):
            return "💰 予算書（金額記入が必要）"
        
        # 報告書系
        if any(kw in fn_lower for kw in ['報告', 'report', '実績']):
            return "📊 報告書フォーマット"
        
        # 事業計画系
        if any(kw in fn_lower for kw in ['計画', '事業', 'plan', 'project']):
            return "📋 事業計画書"
        
        # チェックリスト系
        if any(kw in fn_lower for kw in ['チェック', 'check', '確認', 'リスト']):
            return "✅ チェックリスト"
        
        return "📄 関連資料"
    
    def _classify_file_with_vlm(self, file_path: str, filename: str, grant_name: str = None) -> Optional[str]:
        """
        VLMを使ってファイル内容から種別を判定する。
        助成金名が指定されている場合、その助成金に関連するファイルかも検証する。
        """
        if not self.client:
            return None
        
        try:
            # ファイル内容を抽出
            content = self._extract_file_content_for_classification(file_path)
            if not content:
                return None
            
            # 助成金名の検証を追加 (Positive Matching - 助成金名・交付団体名の含有確認)
            grant_name_check = ""
            if grant_name:
                # 助成金名から主要なキーワードを抽出
                grant_keywords = self._extract_grant_keywords(grant_name)
                grant_name_check = f"""

【最重要】助成金の関連性チェック (Positive Matching):
対象助成金名: {grant_name}
関連キーワード: {', '.join(grant_keywords)}

■ 判定の基本原則（必ず守ること）:
このファイルが「対象助成金」の正規の申請書類であるためには、
ファイル内容に以下のいずれかが**明確に記載されている**必要があります:

1. 助成金名（「{grant_name}」またはその一部）
2. 助成金交付団体名・財団名（上記キーワードのいずれか）
3. 助成金の公募回号や年度と合わせた記載

■ NOT_RELATEDと判定すべきケース:
- ファイル内容に対象助成金名・交付団体名が**一切記載されていない**
- 別の助成金名、別の財団名が明記されている
- 汎用的な書類（会議室申込書、団体設立届、契約書等）
- 財団の活動報告書、パンフレット等の「読み物」

■ 重要: 
「申請書」「様式」などの単語があっても、対象助成金との紐付けがなければNOT_RELATEDです。
"""
            
            prompt = f"""
あなたは助成金申請書類の分類専門家です。
以下のファイルが、対象助成金の正式な申請関連書類であるか厳密に判定してください。
{grant_name_check}

■ ファイル情報:
ファイル名: {filename}

■ ファイル内容（冒頭部分）:
{content[:3000]}

■ 判定手順:
1. まず、ファイル内容に「{grant_name}」またはその交付団体名が明記されているか確認
2. 明記されている場合のみ、以下の分類を行う
3. 明記されていない場合は NOT_RELATED

■ 分類選択肢:
1. APPLICATION_FORM - 対象助成金の申請書/応募書/様式（助成金名が明記されている）
2. GUIDELINES - 対象助成金の募集要項/公募要領（助成金名が明記されている）
3. REGULATIONS - 対象助成金の交付要綱/ガイドライン
4. SAMPLE - 対象助成金の記入例/サンプル
5. BUDGET - 対象助成金の予算書/経費明細様式
6. REPORT - 対象助成金の報告書フォーマット
7. PLAN - 対象助成金の事業計画書様式
8. CHECKLIST - 対象助成金のチェックリスト
9. NOT_RELATED - 対象助成金との関連性なし（助成金名・交付団体名が記載されていない、または別の助成金）
10. OTHER - 対象助成金に関連するがカテゴリ不明

回答は選択肢の英語キー（例: APPLICATION_FORM）のみを出力してください。
"""
            
            response = self.client.models.generate_content(
                model=self.vlm_model,
                contents=prompt
            )
            
            result = response.text.strip().upper()
            
            # 助成金との関連性がない場合はOTHERとして扱う
            if "NOT_RELATED" in result:
                logging.info(f"[FILE_CLASSIFIER] File '{filename}' is not related to grant '{grant_name}'")
                return "📄 関連資料（別の助成金の可能性）"
            
            # 結果をマッピング
            mapping = {
                "APPLICATION_FORM": "📝 申請書フォーマット（記入が必要）",
                "GUIDELINES": "📋 募集要項（応募条件・審査基準が記載）",
                "REGULATIONS": "📜 交付要綱・ガイドライン（ルール・規程）",
                "SAMPLE": "📖 記入例・サンプル（参考資料）",
                "BUDGET": "💰 予算書（金額記入が必要）",
                "REPORT": "📊 報告書フォーマット",
                "PLAN": "📋 事業計画書",
                "CHECKLIST": "✅ チェックリスト",
            }
            
            return mapping.get(result, None)
            
        except Exception as e:
            logging.warning(f"[FILE_CLASSIFIER] VLM classification failed: {e}")
            return None
    
    def _extract_grant_keywords(self, grant_name: str) -> List[str]:
        """
        助成金名から主要なキーワードを抽出する。
        """
        
        keywords = []
        
        # 財団名・法人名の抽出 (例: 「公益財団法人○○財団」→「○○財団」)
        org_patterns = [
            r'(?:公益)?(?:社団|財団)法人\s*([^\s助成]+)',  # 財団法人○○
            r'([^\s]*財団)',  # ○○財団
            r'([^\s]*基金)',  # ○○基金
            r'([^\s]*協会)',  # ○○協会
        ]
        
        for pattern in org_patterns:
            match = re.search(pattern, grant_name)
            if match:
                keywords.append(match.group(1))
        
        # 助成プログラム名の抽出
        program_patterns = [
            r'([^\s]*助成(?:金|プログラム|事業)?)',
            r'([^\s]*支援(?:金|プログラム|事業)?)',
        ]
        
        for pattern in program_patterns:
            match = re.search(pattern, grant_name)
            if match:
                keywords.append(match.group(1))
        
        # 重複を削除しつつ、元の助成金名全体も追加
        if grant_name not in keywords:
            keywords.insert(0, grant_name)
        
        # 重複を削除
        seen = set()
        unique_keywords = []
        for kw in keywords:
            if kw not in seen:
                seen.add(kw)
                unique_keywords.append(kw)
        
        return unique_keywords[:5]  # 最大5つまで
    
    def _extract_file_content_for_classification(self, file_path: str) -> Optional[str]:
        """ファイル内容を抽出（分類用）"""
        try:
            ext = os.path.splitext(file_path)[1].lower()
            
            if ext in ['.docx', '.doc']:
                from docx import Document
                doc = Document(file_path)
                texts = []
                for para in doc.paragraphs[:20]:  # 最初の20段落
                    texts.append(para.text)
                for table in doc.tables[:3]:  # 最初の3テーブル
                    for row in table.rows:
                        texts.append(" | ".join([cell.text.strip() for cell in row.cells]))
                return "\n".join(texts)
            
            elif ext in ['.xlsx', '.xls']:
                import openpyxl
                wb = openpyxl.load_workbook(file_path, read_only=True)
                texts = []
                for sheet_name in wb.sheetnames[:2]:  # 最初の2シート
                    sheet = wb[sheet_name]
                    for row in list(sheet.iter_rows(max_row=20, values_only=True)):
                        row_text = " | ".join([str(cell) for cell in row if cell])
                        if row_text:
                            texts.append(row_text)
                wb.close()
                return "\n".join(texts)
            
            elif ext == '.pdf':
                try:
                    import fitz  # PyMuPDF
                    doc = fitz.open(file_path)
                    texts = []
                    for page_num in range(min(3, doc.page_count)):  # 最初の3ページ
                        page = doc.load_page(page_num)
                        texts.append(page.get_text())
                    doc.close()
                    return "\n".join(texts)
                except ImportError:
                    return None
            
        except Exception as e:
            logging.warning(f"[FILE_CLASSIFIER] Content extraction failed: {e}")
        
        return None
