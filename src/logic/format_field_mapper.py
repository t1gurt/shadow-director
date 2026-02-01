"""
Format Field Mapper - VLMを使用してExcel/Wordフォーマットのフィールドを検出し、
ドラフト内容とマッピングする。

Cloud Run（Linux）対応。
"""

import logging
import json
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
import os

from google import genai
from google.genai import types
from src.utils.office_utils import convert_to_pdf


@dataclass
class FieldInfo:
    """フォーマット内のフィールド情報"""
    field_id: str           # 一意のID
    field_name: str         # フィールド名（ラベル）
    field_type: str         # "cell" | "table_cell" | "paragraph" | "text_box"
    location: Dict[str, Any]  # 位置情報（Excel: {"sheet": str, "row": int, "col": int}, Word: {"paragraph_idx": int} or {"table_idx": int, "row": int, "col": int}）
    max_length: Optional[int] = None  # 文字数制限
    description: Optional[str] = None  # 説明
    required: bool = False  # 必須項目
    input_length_type: str = "unknown"  # "short"（1行以内）| "long"（100字以上の長文）| "unknown"


class FormatFieldMapper:
    """
    VLMでExcel/Wordフォーマットのフィールドを検出し、
    ドラフト内容とマッピングする。
    """
    
    def __init__(self, gemini_client=None, model_name: str = "gemini-3.0-pro"):
        """
        Args:
            gemini_client: Gemini APIクライアント
            model_name: 使用するモデル名
        """
        self.client = gemini_client
        self.model_name = model_name
        self.logger = logging.getLogger(__name__)
        
        # フィールド制限の情報を保持
        self.last_skipped_field_count = 0
        self.last_total_field_count = 0
        
        if not self.client:
            try:
                from src.utils.gemini_client import get_gemini_client
                self.client = get_gemini_client()
            except Exception as e:
                self.logger.error(f"[FORMAT_MAPPER] Failed to initialize Gemini client: {e}")
    
    def analyze_excel_fields(self, file_path: str) -> List[FieldInfo]:
        """
        Excelファイルのフィールドを解析する。
        
        Args:
            file_path: Excelファイルのパス
            
        Returns:
            検出されたフィールドのリスト
        """
        fields = []
        
        try:
            import openpyxl
            
            wb = openpyxl.load_workbook(file_path, read_only=True)
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_fields = self._analyze_excel_sheet(sheet, sheet_name)
                fields.extend(sheet_fields)
            
            wb.close()
            
            self.logger.info(f"[FORMAT_MAPPER] Found {len(fields)} fields in Excel file")
            
            # VLMで追加解析（フィールドの意味を推論）
            if fields and self.client:
                fields = self._enhance_fields_with_vlm(fields, file_path, "excel")
            
        except ImportError:
            self.logger.warning("[FORMAT_MAPPER] openpyxl not installed")
        except Exception as e:
            self.logger.error(f"[FORMAT_MAPPER] Excel analysis failed: {e}")
        
        return fields
    
    def _analyze_excel_sheet(self, sheet, sheet_name: str) -> List[FieldInfo]:
        """
        Excelシートから入力フィールドを検出する。
        
        入力欄のパターン:
        - 空のセル（隣にラベルあり）
        - 下線（_____）で示された入力欄
        - 黄色/水色などの背景色のセル
        """
        fields = []
        
        try:
            # シートの使用範囲を取得
            max_row = min(sheet.max_row or 100, 100)
            max_col = min(sheet.max_column or 20, 20)
            
            for row in range(1, max_row + 1):
                for col in range(1, max_col + 1):
                    cell = sheet.cell(row=row, column=col)
                    
                    # ラベルセルを探す（テキストが入っている）
                    if cell.value and isinstance(cell.value, str):
                        cell_text = str(cell.value).strip()
                        
                        # 入力欄を示すキーワード
                        if any(keyword in cell_text for keyword in ['記入', '入力', '：', ':', '）', ')']):
                            # 右隣または下のセルが入力欄
                            right_cell = sheet.cell(row=row, column=col + 1)
                            below_cell = sheet.cell(row=row + 1, column=col) if row < max_row else None
                            
                            # 右隣が空なら入力欄候補
                            if right_cell and (right_cell.value is None or right_cell.value == ''):
                                field = FieldInfo(
                                    field_id=f"{sheet_name}_{row}_{col+1}",
                                    field_name=cell_text,
                                    field_type="cell",
                                    location={
                                        "sheet": sheet_name,
                                        "row": row,
                                        "col": col + 1
                                    }
                                )
                                fields.append(field)
                            # 下が空なら入力欄候補
                            elif below_cell and (below_cell.value is None or below_cell.value == ''):
                                field = FieldInfo(
                                    field_id=f"{sheet_name}_{row+1}_{col}",
                                    field_name=cell_text,
                                    field_type="cell",
                                    location={
                                        "sheet": sheet_name,
                                        "row": row + 1,
                                        "col": col
                                    }
                                )
                                fields.append(field)
                        
                        # 文字数制限を検出（例：「（400字以内）」）
                        import re
                        limit_match = re.search(r'(\d+)\s*[字文]', cell_text)
                        if limit_match:
                            # この近くのフィールドに制限を適用
                            for field in fields[-3:]:  # 直近3フィールド
                                if field.max_length is None:
                                    field.max_length = int(limit_match.group(1))
                                    break
        
        except Exception as e:
            self.logger.warning(f"[FORMAT_MAPPER] Error analyzing sheet {sheet_name}: {e}")
        
        return fields
    


    def _analyze_word_with_pdf_vlm(self, file_path: str) -> List[FieldInfo]:
        """PDF変換を介してWordドキュメントをVLM解析する（.doc対応）"""
        fields = []
        try:
            # 1. Word -> PDF変換
            self.logger.info(f"[FORMAT_MAPPER] Converting Word to PDF for VLM: {file_path}")
            pdf_path = convert_to_pdf(file_path)
            
            if not pdf_path:
                self.logger.warning("[FORMAT_MAPPER] PDF conversion failed")
                return fields
                
            # 2. PDFをVLMに送信
            import pathlib
            pdf_data = pathlib.Path(pdf_path).read_bytes()
            
            prompt = """
この申請書の画像から、記入が必要なフィールド（入力欄）を全て抽出してください。
以下の情報をJSON形式で出力してください:
[
  {
    "field_id": "ユニークなID",
    "field_name": "項目名",
    "description": "入力内容の説明や制約",
    "max_length": "文字数制限（あれば数値、なければnull）",
    "input_length_type": "short または long"
  }
]
"""
            contents = [
                types.Part(
                    inline_data=types.Blob(
                        mime_type="application/pdf",
                        data=pdf_data
                    )
                ),
                types.Part(text=prompt)
            ]
            
            response = self.client.models.generate_content(
                model=self.model_name,
                contents=contents,
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    max_output_tokens=8192,
                    temperature=0.1
                )
            )
            
            # JSONパース
            response_text = response.text.strip()
            # Markdownコードブロック除去
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0]
            elif "```" in response_text:
                response_text = response_text.split("```")[1].split("```")[0]
                
            data = json.loads(response_text)
            
            if isinstance(data, list):
                for i, item in enumerate(data):
                    field = FieldInfo(
                        field_id=item.get("field_id", f"vlm_pdf_{i}"),
                        field_name=item.get("field_name", "不明"),
                        field_type="vlm_detected",
                        description=item.get("description"),
                        max_length=item.get("max_length"),
                        input_length_type=item.get("input_length_type", "short"),
                        location={"source": "pdf_vlm", "index": i}
                    )
                    fields.append(field)
            
            self.logger.info(f"[FORMAT_MAPPER] PDF VLM found {len(fields)} fields")
            
            # クリーンアップ（PDF削除）
            try:
                os.remove(pdf_path)
            except:
                pass
                
        except Exception as e:
            self.logger.error(f"[FORMAT_MAPPER] PDF VLM analysis failed: {e}")
            
        return fields

    def analyze_word_fields(self, file_path: str) -> List[FieldInfo]:
        """
        Wordファイルのフィールドを解析する。
        """
        fields = []
        
        # .doc (バイナリ形式) の場合は直接PDF変換VLMへ
        if file_path.lower().endswith('.doc'):
            self.logger.info(f"[FORMAT_MAPPER] Detected .doc file, using PDF VLM analysis: {file_path}")
            if self.client:
                return self._analyze_word_with_pdf_vlm(file_path)
            else:
                self.logger.warning("[FORMAT_MAPPER] .doc file detected but VLM client is not active")
                return []
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # ドキュメント要素（段落とテーブル）の順序を構築（行コンテキスト判定用）
            block_items = []
            try:
                # python-docxの内部要素にアクセスして順序を特定
                from docx.oxml.text.paragraph import CT_P
                from docx.oxml.table import CT_Tbl
                
                p_iter = iter(doc.paragraphs)
                t_iter = iter(doc.tables)
                p_idx = 0
                t_idx = 0
                
                # body内の要素を順に走査
                for child in doc.element.body.iterchildren():
                    if isinstance(child, CT_P):
                        try:
                            block_items.append({"type": "paragraph", "obj": next(p_iter), "index": p_idx})
                            p_idx += 1
                        except StopIteration:
                            break
                    elif isinstance(child, CT_Tbl):
                        try:
                            block_items.append({"type": "table", "obj": next(t_iter), "index": t_idx})
                            t_idx += 1
                        except StopIteration:
                            break
            except Exception as e:
                self.logger.warning(f"[FORMAT_MAPPER] Failed to build block items logic: {e}")
                # フォールバック: 段落のみのリスト作成
                block_items = [{"type": "paragraph", "obj": p, "index": i} for i, p in enumerate(doc.paragraphs)]

            # テーブルからフィールドを検出
            for table_idx, table in enumerate(doc.tables):
                table_fields = self._analyze_word_table(table, table_idx, block_items)
                fields.extend(table_fields)
            
            # 段落からフィールドを検出（順序情報を渡す）
            paragraph_fields = self._analyze_word_paragraphs(doc.paragraphs, block_items)
            fields.extend(paragraph_fields)
            
            self.logger.info(f"[FORMAT_MAPPER] Found {len(fields)} fields in Word file (text-based)")
            
            # テキストベースで検出できなかった場合、VLMでドキュメント全体を解析
            if not fields and self.client:
                self.logger.info("[FORMAT_MAPPER] No fields found with text-based analysis, trying VLM (PDF-based)...")
                # PDF変換VLMを優先
                fields = self._analyze_word_with_pdf_vlm(file_path)
                
                # 失敗時は従来のテキストベースVLMへ
                if not fields:
                     self.logger.info("[FORMAT_MAPPER] PDF VLM returned no fields, falling back to text-based VLM...")
                     fields = self._analyze_word_with_vlm(doc, file_path)
            
            # VLMで追加解析（フィールドの意味を推論）
            if fields and self.client:
                fields = self._enhance_fields_with_vlm(fields, file_path, "word")
            
        except ImportError:
            self.logger.warning("[FORMAT_MAPPER] python-docx not installed")
        except Exception as e:
            self.logger.error(f"[FORMAT_MAPPER] Word analysis failed: {e}")
            # エラー時（.docなど）のフォールバック
            if self.client and not fields:
                self.logger.info("[FORMAT_MAPPER] Attempting PDF VLM fallback after error...")
                fields = self._analyze_word_with_pdf_vlm(file_path)
        
        return fields
    
    def _fallback_word_all_cells(self, doc) -> List[FieldInfo]:
        """
        フォールバック: すべてのテーブルセルを入力候補として検出。
        ラベルセルの隣の空セルまたは長いセルを入力欄として扱う。
        """
        fields = []
        
        try:
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    cells = row.cells
                    for col_idx, cell in enumerate(cells):
                        cell_text = cell.text.strip()
                        
                        # 空セル（入力欄候補）
                        if not cell_text:
                            # 左隣にラベルがあれば入力欄
                            if col_idx > 0:
                                label_cell = cells[col_idx - 1]
                                label_text = label_cell.text.strip()
                                if label_text and len(label_text) < 50:
                                    field = FieldInfo(
                                        field_id=f"table{table_idx}_{row_idx}_{col_idx}",
                                        field_name=label_text,
                                        field_type="table_cell",
                                        location={
                                            "table_idx": table_idx,
                                            "row": row_idx,
                                            "col": col_idx
                                        }
                                    )
                                    fields.append(field)
                        
                        # 長いセル（説明や記述欄の可能性）
                        elif len(cell_text) > 50 and col_idx > 0:
                            # 左隣がラベルなら入力欄
                            label_cell = cells[col_idx - 1]
                            label_text = label_cell.text.strip()
                            if label_text and len(label_text) < 50:
                                field = FieldInfo(
                                    field_id=f"table{table_idx}_{row_idx}_{col_idx}",
                                    field_name=label_text,
                                    field_type="table_cell",
                                    location={
                                        "table_idx": table_idx,
                                        "row": row_idx,
                                        "col": col_idx
                                    }
                                )
                                fields.append(field)
            
            self.logger.info(f"[FORMAT_MAPPER] Fallback detected {len(fields)} fields")
            
        except Exception as e:
            self.logger.warning(f"[FORMAT_MAPPER] Fallback analysis failed: {e}")
        
        return fields
    
    def _analyze_word_with_vlm(self, doc, file_path: str) -> List[FieldInfo]:
        """VLMを使ってWordドキュメント全体を解析しフィールドを抽出（チャンク分割対応）"""
        all_fields = []
        
        try:
            # チャンクサイズの設定（1チャンクあたり最大30フィールドを目安）
            CHUNK_SIZE = 4000  # 文字数ベースでチャンク分割
            
            # ドキュメントの全テキストを構造を保って抽出
            chunks = []
            current_chunk = ""
            chunk_start_table = 0
            chunk_start_para = 0
            
            # 段落を収集
            para_texts = []
            for para_idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    para_texts.append((para_idx, text))
            
            # テーブルを収集
            table_texts = []
            for table_idx, table in enumerate(doc.tables):
                table_text = f"\n[テーブル{table_idx}]\n"
                for row_idx, row in enumerate(table.rows):
                    row_text = " | ".join([f"[{col_idx}]{cell.text.strip()}" for col_idx, cell in enumerate(row.cells)])
                    table_text += f"  Row{row_idx}: {row_text}\n"
                table_texts.append((table_idx, table_text))
            
            # 段落をチャンクに分割
            para_chunk = ""
            para_chunk_indices = []
            for para_idx, text in para_texts:
                line = f"[段落{para_idx}] {text}\n"
                if len(para_chunk) + len(line) > CHUNK_SIZE:
                    if para_chunk:
                        chunks.append(("paragraph", para_chunk, para_chunk_indices.copy()))
                        para_chunk = ""
                        para_chunk_indices = []
                para_chunk += line
                para_chunk_indices.append(para_idx)
            if para_chunk:
                chunks.append(("paragraph", para_chunk, para_chunk_indices))
            
            # テーブルをチャンクに分割
            for table_idx, table_text in table_texts:
                if len(table_text) > CHUNK_SIZE:
                    # 大きなテーブルはさらに分割
                    lines = table_text.split('\n')
                    sub_chunk = ""
                    for line in lines:
                        if len(sub_chunk) + len(line) > CHUNK_SIZE:
                            if sub_chunk:
                                chunks.append(("table", sub_chunk, [table_idx]))
                                sub_chunk = ""
                        sub_chunk += line + "\n"
                    if sub_chunk:
                        chunks.append(("table", sub_chunk, [table_idx]))
                else:
                    chunks.append(("table", table_text, [table_idx]))
            
            self.logger.info(f"[FORMAT_MAPPER] Split document into {len(chunks)} chunks for VLM analysis")
            
            # 各チャンクに対してVLM解析を実行
            for chunk_idx, (chunk_type, chunk_text, indices) in enumerate(chunks):
                try:
                    chunk_fields = self._analyze_chunk_with_vlm(chunk_text, chunk_type, chunk_idx, len(chunks))
                    all_fields.extend(chunk_fields)
                    self.logger.info(f"[FORMAT_MAPPER] Chunk {chunk_idx + 1}/{len(chunks)}: found {len(chunk_fields)} fields")
                except Exception as e:
                    self.logger.warning(f"[FORMAT_MAPPER] Chunk {chunk_idx + 1} analysis failed: {e}")
            
            # 重複フィールドを除去（field_idベース）
            seen_ids = set()
            unique_fields = []
            for field in all_fields:
                if field.field_id not in seen_ids:
                    seen_ids.add(field.field_id)
                    unique_fields.append(field)
            
            self.logger.info(f"[FORMAT_MAPPER] VLM detected total {len(unique_fields)} unique fields (from {len(all_fields)} raw)")
            
            # 検出されたパターンをログに出力
            pattern_counts = {}
            for f in unique_fields:
                pattern = f.location.get("input_pattern", "unknown")
                pattern_counts[pattern] = pattern_counts.get(pattern, 0) + 1
            self.logger.info(f"[FORMAT_MAPPER] Pattern breakdown: {pattern_counts}")
            
            return unique_fields
            
        except Exception as e:
            self.logger.error(f"[FORMAT_MAPPER] VLM Word analysis failed: {e}")
        
        return all_fields
    
    def _analyze_chunk_with_vlm(self, chunk_text: str, chunk_type: str, chunk_idx: int, total_chunks: int) -> List[FieldInfo]:
        """単一チャンクをVLMで解析してフィールドを抽出"""
        fields = []
        
        if not self.client or not chunk_text.strip():
            return fields
        
        # VLMにフィールド検出を依頼（入力パターン検出を追加）
        prompt = f"""
以下はWord申請書フォーマットの一部（チャンク {chunk_idx + 1}/{total_chunks}）です。
入力が必要なフィールドを検出し、**入力パターン**を特定してください。

## ドキュメント内容
{chunk_text}

## 入力パターンの種類
1. **inline** - ラベルと同一行、コロン後に入力する（例：「団体名：」の後に入力）
2. **next_line** - ラベルの次の行に入力する（例：「1. 事業概要（200字以内）」の次行に入力）
3. **underline** - 下線プレースホルダー「____」を置換する
4. **bracket** - 括弧プレースホルダー「（　）」を置換する
5. **table_cell** - テーブルセル内に入力する（ラベルセルの隣の空セル）

## 出力形式（JSON配列）
[
  {{
    "field_id": "para_5 または table0_1_2",
    "field_name": "ラベル名（例：団体名、事業概要）",
    "field_type": "paragraph または table_cell",
    "input_pattern": "inline | next_line | underline | bracket | table_cell のいずれか",
    "table_idx": テーブル番号（0から、table_cellの場合のみ）,
    "row": 行番号（table_cellの場合のみ）,
    "col": 列番号（入力するセルの列、table_cellの場合のみ）,
    "paragraph_idx": 段落番号（paragraphの場合、入力する段落の番号）,
    "label_paragraph_idx": ラベルがある段落番号（next_lineの場合のみ）,
    "description": "入力すべき内容の説明",
    "max_length": 文字数制限があれば数値、なければnull
  }}
]

## 重要な注意
- **paragraph_idx**は実際に入力を行う段落の番号です
- next_lineパターンの場合、paragraph_idxはラベルの次の段落です
- inlineパターンの場合、paragraph_idxはラベルと同じ段落です
- 空欄、下線「____」、括弧「（　）」、「記入してください」などの指示がある箇所を検出
- フィールドがない場合は空配列[]を返してください
- JSONのみを出力してください
"""
        
        response = self.client.models.generate_content(
            model=self.model_name,
            contents=prompt
        )
        
        # JSONをパース
        response_text = response.text.strip()
        if "```json" in response_text:
            response_text = response_text.split("```json")[1].split("```")[0]
        elif "```" in response_text:
            response_text = response_text.split("```")[1].split("```")[0]
        
        vlm_fields = json.loads(response_text)
        
        for vf in vlm_fields:
            input_pattern = vf.get("input_pattern", "inline")
            
            if vf.get("field_type") == "table_cell" or input_pattern == "table_cell":
                location = {
                    "table_idx": vf.get("table_idx", 0),
                    "row": vf.get("row", 0),
                    "col": vf.get("col", 0),
                    "input_pattern": "table_cell"
                }
                field_type = "table_cell"
            else:
                location = {
                    "paragraph_idx": vf.get("paragraph_idx", 0),
                    "input_pattern": input_pattern,
                    "label_paragraph_idx": vf.get("label_paragraph_idx")
                }
                field_type = "paragraph"
            
            field = FieldInfo(
                field_id=vf.get("field_id", f"vlm_chunk{chunk_idx}_{len(fields)}"),
                field_name=vf.get("field_name", ""),
                field_type=field_type,
                location=location,
                description=vf.get("description"),
                max_length=vf.get("max_length")
            )
            fields.append(field)
        
        return fields
    
    def _analyze_word_table(self, table, table_idx: int, block_items=None) -> List[FieldInfo]:
        """
        Wordテーブルからフィールドを検出（複数パターン対応）
        
        対応パターン:
        1. 1セル(1行1列)のテーブル（新規対応）
        2. ラベル | 入力欄（空） のペア
        3. 1行に複数の ラベル | 入力欄 ペアがある場合
        4. ヘッダー行 + データ行パターン（例：｜項目｜金額｜説明｜）
        """
        fields = []
        
        try:
            rows = list(table.rows)
            if not rows:
                return fields
            
            # 1セルテーブルの場合は専用ロジック
            if len(rows) == 1 and len(rows[0].cells) == 1:
                self.logger.info(f"[FORMAT_MAPPER] Table {table_idx}: detected single-cell table")
                return self._detect_single_cell_pattern(table, table_idx, block_items)
            
            # まず既存のラベル→入力パターンで検出を試みる
            label_input_fields = self._detect_label_input_pattern(table, table_idx, rows)
            
            # ヘッダー行パターンも検出
            header_row_fields = self._detect_header_row_pattern(table, table_idx, rows)
            
            # どちらか多く検出できた方を採用（両方併用すると重複の可能性）
            if len(header_row_fields) > len(label_input_fields):
                fields = header_row_fields
                self.logger.info(f"[FORMAT_MAPPER] Table {table_idx}: using header-row pattern, found {len(fields)} fields")
            else:
                fields = label_input_fields
                self.logger.info(f"[FORMAT_MAPPER] Table {table_idx}: using label-input pattern, found {len(fields)} fields")
        
        except Exception as e:
            self.logger.warning(f"[FORMAT_MAPPER] Error analyzing Word table: {e}")
        
        return fields
    
    def _detect_single_cell_pattern(self, table, table_idx: int, block_items=None) -> List[FieldInfo]:
        """
        1セル(1行1列)のテーブルを検出する。
        
        対応パターン:
        1. 段落ラベル + 1セルテーブル(入力欄)
           例: <メールアドレス> (段落)
               ┌────────┐
               │        │ (1セルテーブル)
               └────────┘
        
        2. 1セル内にラベル + プレースホルダー
           例: ┌──────────────┐
               │ ラベル：____ │
               └──────────────┘
        
        Args:
            table: Wordのテーブルオブジェクト
            table_idx: テーブルのインデックス
            block_items: ドキュメント要素の順序リスト（任意）
        
        Returns:
            検出されたフィールドのリスト
        """
        fields = []
        
        try:
            cell = table.rows[0].cells[0]
            cell_text = cell.text.strip()
            
            # パターン2: 1セル内にラベル + プレースホルダー
            # 下線プレースホルダー: "ラベル：_____"
            # 括弧プレースホルダー: "ラベル：（　）"
            import re
            
            # 下線パターン
            underline_match = re.search(r'(.+?)[:：]\s*[_＿]{3,}', cell_text)
            if underline_match:
                field = FieldInfo(
                    field_id=f"table{table_idx}_0_0",
                    field_name=underline_match.group(1).strip(),
                    field_type="table_cell",
                    location={
                        "table_idx": table_idx,
                        "row": 0,
                        "col": 0,
                        "input_pattern": "single_cell_underline"
                    }
                )
                fields.append(field)
                self.logger.info(f"[FORMAT_MAPPER] Single-cell table {table_idx}: detected underline pattern")
                return fields
            
            # 括弧パターン
            bracket_match = re.search(r'(.+?)[(（]\s*[　\s]*[)）]', cell_text)
            if bracket_match:
                field = FieldInfo(
                    field_id=f"table{table_idx}_0_0",
                    field_name=bracket_match.group(1).strip(),
                    field_type="table_cell",
                    location={
                        "table_idx": table_idx,
                        "row": 0,
                        "col": 0,
                        "input_pattern": "single_cell_bracket"
                    }
                )
                fields.append(field)
                self.logger.info(f"[FORMAT_MAPPER] Single-cell table {table_idx}: detected bracket pattern")
                return fields
            
            # パターン1: 空セルまたはプレースホルダーのみ → 直前の段落がラベル
            is_empty = not cell_text
            is_placeholder = cell_text and (
                all(c in '_＿　 ' for c in cell_text) or
                cell_text in ['（　）', '(　)', '（）', '()']
            )
            
            if is_empty or is_placeholder:
                # 直前の段落をラベルとして検索
                label = self._find_label_before_table(table_idx, block_items)
                
                if label:
                    field = FieldInfo(
                        field_id=f"table{table_idx}_0_0",
                        field_name=label,
                        field_type="table_cell",
                        location={
                            "table_idx": table_idx,
                            "row": 0,
                            "col": 0,
                            "input_pattern": "single_cell_with_paragraph_label"
                        }
                    )
                    fields.append(field)
                    self.logger.info(f"[FORMAT_MAPPER] Single-cell table {table_idx}: detected with paragraph label '{label}'")
                else:
                    # ラベルが見つからない場合でも、空の1セルテーブルとして検出
                    field = FieldInfo(
                        field_id=f"table{table_idx}_0_0",
                        field_name=f"テーブル{table_idx}",
                        field_type="table_cell",
                        location={
                            "table_idx": table_idx,
                            "row": 0,
                            "col": 0,
                            "input_pattern": "single_cell_no_label"
                        }
                    )
                    fields.append(field)
                    self.logger.info(f"[FORMAT_MAPPER] Single-cell table {table_idx}: detected without label")
        
        except Exception as e:
            self.logger.warning(f"[FORMAT_MAPPER] Error detecting single-cell table: {e}")
        
        return fields
    
    def _find_label_before_table(self, table_idx: int, block_items=None) -> Optional[str]:
        """
        テーブルの直前の段落からラベルを取得する。
        
        Args:
            table_idx: テーブルのインデックス
            block_items: ドキュメント要素の順序リスト
        
        Returns:
            ラベル文字列、見つからない場合はNone
        """
        if not block_items:
            return None
        
        try:
            # block_items内のテーブル要素を探す
            for idx, item in enumerate(block_items):
                if item["type"] == "table" and item["index"] == table_idx:
                    # 直前の要素を確認
                    if idx > 0:
                        prev_item = block_items[idx - 1]
                        if prev_item["type"] == "paragraph":
                            # 段落のテキストを取得
                            para_text = prev_item["obj"].text.strip()
                            
                            # ラベルとして適切か判定（短く、特定のパターンに一致）
                            # 例: "<メールアドレス>", "団体名：", "1. 団体登録ID"
                            if para_text and len(para_text) < 100:
                                # 末尾のコロンや記号を除去
                                import re
                                label = re.sub(r'[:：]+\s*$', '', para_text)
                                # 番号付きラベルの場合、番号を含める
                                return label
            
        except Exception as e:
            self.logger.warning(f"[FORMAT_MAPPER] Error finding label before table: {e}")
        
        return None
    
    def _detect_label_input_pattern(self, table, table_idx: int, rows) -> List[FieldInfo]:
        """ラベル→入力欄のペアパターンを検出"""
        fields = []
        
        for row_idx, row in enumerate(rows):
            cells = row.cells
            num_cells = len(cells)
            
            # 処理済みセルを追跡（1行内で複数フィールドを検出するため）
            processed_cols = set()
            
            col_idx = 0
            while col_idx < num_cells:
                # 既に処理済みのセルはスキップ
                if col_idx in processed_cols:
                    col_idx += 1
                    continue
                
                cell = cells[col_idx]
                cell_text = cell.text.strip()
                
                # ラベルセルを探す（テキストがあり、短めのもの）
                if cell_text and len(cell_text) < 100:
                    # 次のセルが空または入力欄候補なら入力欄
                    if col_idx + 1 < num_cells:
                        next_cell = cells[col_idx + 1]
                        next_text = next_cell.text.strip()
                        
                        # 空セル、下線、または空白のみの場合は入力欄
                        is_empty = not next_text
                        is_placeholder = next_text and (
                            all(c in '_＿　 ' for c in next_text) or
                            next_text in ['（　）', '(　)', '（）', '()']
                        )
                        
                        if is_empty or is_placeholder:
                            field = FieldInfo(
                                field_id=f"table{table_idx}_{row_idx}_{col_idx+1}",
                                field_name=cell_text,
                                field_type="table_cell",
                                location={
                                    "table_idx": table_idx,
                                    "row": row_idx,
                                    "col": col_idx + 1,
                                    "input_pattern": "table_cell"
                                }
                            )
                            fields.append(field)
                            
                            # ラベルセルと入力セルの両方を処理済みにマーク
                            processed_cols.add(col_idx)
                            processed_cols.add(col_idx + 1)
                            
                            # 次のセルへスキップ（入力欄の次から再開）
                            col_idx += 2
                            continue
                
                col_idx += 1
        
        return fields
    
    def _detect_header_row_pattern(self, table, table_idx: int, rows) -> List[FieldInfo]:
        """
        ヘッダー行 + データ行パターンを検出
        例: | 項目 | 金額 | 説明 | コメント |
            | xxx  | xxx  | xxx  | xxx      |
        """
        fields = []
        
        if len(rows) < 2:
            return fields
        
        # 1行目をヘッダー候補として検査
        first_row = rows[0]
        header_cells = first_row.cells
        
        # ヘッダー行の条件: 全てのセルにテキストがある、かつ短いテキスト
        header_texts = [cell.text.strip() for cell in header_cells]
        
        # ヘッダーらしさを判定
        if not all(text and len(text) < 50 for text in header_texts):
            return fields
        
        # 2行目以降のセルが入力欄候補かどうかを確認
        data_rows_count = 0
        for row_idx in range(1, len(rows)):
            row = rows[row_idx]
            cells = row.cells
            
            # データ行の条件: 少なくとも1つのセルが空または入力欄っぽい
            empty_or_placeholder_count = 0
            for cell in cells:
                text = cell.text.strip()
                is_empty = not text
                is_placeholder = text and (
                    all(c in '_＿　 ' for c in text) or
                    text in ['（　）', '(　)', '（）', '()']
                )
                if is_empty or is_placeholder:
                    empty_or_placeholder_count += 1
            
            # データ行と判定（半分以上が空または入力欄候補）
            if empty_or_placeholder_count >= len(cells) // 2:
                data_rows_count += 1
        
        # データ行が1つ以上あればヘッダー行パターンと判定
        if data_rows_count == 0:
            return fields
        
        # ヘッダー行パターンが確定 → 各データ行の各セルをフィールドとして登録
        
        # まず、入力対象外とする列を判定する
        # その列のデータ行いずれかに値（プレースホルダー以外）が入っている場合、その列は「計算式」や「記入済み」とみなして除外
        excluded_col_indices = set()
        for col_idx in range(len(header_texts)):
            is_column_prefilled = False
            for row_idx in range(1, len(rows)):
                if col_idx < len(rows[row_idx].cells):
                    cell_text = rows[row_idx].cells[col_idx].text.strip()
                    # 空でもプレースホルダーでもない場合、何らかの値が入っている
                    is_empty = not cell_text
                    is_placeholder = cell_text and (
                        all(c in '_＿　 ' for c in cell_text) or
                        cell_text in ['（　）', '(　)', '（）', '()']
                    )
                    if not is_empty and not is_placeholder:
                        is_column_prefilled = True
                        break
            
            if is_column_prefilled:
                excluded_col_indices.add(col_idx)
        
        for row_idx in range(1, len(rows)):
            row = rows[row_idx]
            cells = row.cells
            
            for col_idx, cell in enumerate(cells):
                # 入力対象外の列はスキップ
                if col_idx in excluded_col_indices:
                    continue
                
                # ヘッダーのラベルを取得
                if col_idx < len(header_texts):
                    header_label = header_texts[col_idx]
                else:
                    header_label = f"列{col_idx + 1}"
                
                # 既に値が入っているセルはスキップ（入力欄として認識しない）
                # ※列除外ロジックを入れたので、ここは念の為の二重チェック
                cell_text = cell.text.strip()
                is_empty = not cell_text
                is_placeholder = cell_text and (
                    all(c in '_＿　 ' for c in cell_text) or
                    cell_text in ['（　）', '(　)', '（）', '()']
                )
                
                if is_empty or is_placeholder:
                    # 行コンテキストの決定（左端列の値を使用するか、行番号を使用するか）
                    row_context = f"行{row_idx}"
                    # 左端のセル（列0）にテキストがあれば、それを行ラベルとして使用する
                    if len(cells) > 0:
                        first_cell_text = cells[0].text.strip()
                        # テキストがあり、かつ長すぎない（ラベルとして成立する）場合
                        if first_cell_text and len(first_cell_text) < 50:
                            # 括弧や改行を整理
                            row_context = first_cell_text.replace('\n', ' ').strip()
                    
                    # 行番号または行ラベルをラベルに含める
                    field_name = f"{header_label}（{row_context}）" if len(rows) > 2 else header_label
                    
                    field = FieldInfo(
                        field_id=f"table{table_idx}_{row_idx}_{col_idx}",
                        field_name=field_name,
                        field_type="table_cell",
                        location={
                            "table_idx": table_idx,
                            "row": row_idx,
                            "col": col_idx,
                            "input_pattern": "table_cell",
                            "header_label": header_label
                        }
                    )
                    fields.append(field)
        
        return fields
    
    def _analyze_word_paragraphs(self, paragraphs, block_items=None) -> List[FieldInfo]:
        """
        Word段落からフィールドを検出。
        
        Args:
            paragraphs: 段落オブジェクトのリスト
            block_items: ドキュメント要素の順序リスト（任意）
            
        対応パターン:
        - ラベル：____（下線プレースホルダー）
        - ラベル：（入力してください）
        - ラベル：　　（空白プレースホルダー）
        - ◻ ラベル：（チェックボックス形式）
        - 1. 質問文（200字以内）→ 次行入力
        - ラベル：（同一行に短い入力がある場合）
        - 入れ子質問: 4. 親質問 → ①サブ項目（文字数制限）
        """
        fields = []
        
        # block_itemsが渡されていない場合の簡易マッピング
        para_to_block_idx = {}
        if block_items:
            for idx, item in enumerate(block_items):
                if item["type"] == "paragraph":
                    para_to_block_idx[item["index"]] = idx
        
        try:
            import re
            
            # 親質問のコンテキストを追跡
            parent_question = None
            parent_question_idx = None
            
            for para_idx, para in enumerate(paragraphs):
                text = para.text.strip()
                
                if not text or len(text) < 3:
                    continue
                
                # 次の要素がテーブルかどうかを確認するヘルパー
                def is_next_element_table(current_para_idx):
                    if not block_items or current_para_idx not in para_to_block_idx:
                        return False
                    
                    current_block_idx = para_to_block_idx[current_para_idx]
                    if current_block_idx + 1 < len(block_items):
                        next_item = block_items[current_block_idx + 1]
                        if next_item["type"] == "table":
                            return True
                    return False

                # 入れ子パターン: サブ項目を検出 ①②③, (1)(2)(3), ア.イ.ウ. 等
                sub_item_match = re.match(r'^[　\s]*([①②③④⑤⑥⑦⑧⑨⑩]|[(（][1-9１-９][)）]|[ア-オ][.．])(.+?)([（(]\d+字?以?内?[)）])?$', text)
                if sub_item_match and parent_question:
                    sub_marker = sub_item_match.group(1)
                    sub_label = sub_item_match.group(2).strip()
                    char_limit = sub_item_match.group(3)
                    
                    # 親質問名とサブ項目名を組み合わせる
                    combined_label = f"{parent_question} - {sub_marker}{sub_label}"
                    
                    # 次がテーブルなら、この段落は入力欄ではない（テーブルが入力欄）
                    if is_next_element_table(para_idx):
                        continue

                    field = FieldInfo(
                        field_id=f"para_{para_idx + 1}",  # 次の段落が入力欄
                        field_name=combined_label,
                        field_type="paragraph",
                        location={
                            "paragraph_idx": para_idx + 1,
                            "input_type": "next_line",
                            "parent_question": parent_question,
                            "parent_idx": parent_question_idx,
                            "sub_marker": sub_marker
                        },
                        description=char_limit,
                        max_length=int(re.search(r'\d+', char_limit).group()) if char_limit and re.search(r'\d+', char_limit) else None
                    )
                    fields.append(field)
                    continue
                
                # パターン1: チェックボックス形式 "◻ ラベル：入力"
                checkbox_match = re.match(r'^[◻☐□■◼◾▢☑✓✔]?\s*(.+?)[:：]\s*(.*)$', text)
                if checkbox_match:
                    label = checkbox_match.group(1).strip()
                    value_hint = checkbox_match.group(2).strip()
                    
                    # 括弧付きヒント or 空欄の場合は入力欄
                    if not value_hint or re.match(r'^[（(].+[)）]$', value_hint) or re.match(r'^[　\s]+$', value_hint) or re.match(r'^[_＿]+$', value_hint):
                        field = FieldInfo(
                            field_id=f"para_{para_idx}",
                            field_name=label,
                            field_type="paragraph",
                            location={
                                "paragraph_idx": para_idx,
                                "input_type": "inline"  # 同一行入力
                            },
                            description=value_hint if value_hint else None
                        )
                        fields.append(field)
                        continue
                
                # パターン2: 番号付き質問 "1. 質問文（文字数）" → 次の段落が入力欄
                # この質問が「親質問」になる可能性がある
                question_match = re.match(r'^(\d+)[.．、]\s*(.+?)([（(]\d+字?以?内?[)）])?$', text)
                if question_match:
                    label = question_match.group(2).strip()
                    char_limit = question_match.group(3)
                    
                    # この質問を親質問として記憶（後続のサブ項目のため）
                    parent_question = f"{question_match.group(1)}. {label}"
                    parent_question_idx = para_idx
                    
                    # 文字数制限がある場合は直接の入力欄
                    # ただし、次がテーブルなら入力をスキップ（テーブル側で処理）
                    if char_limit and not is_next_element_table(para_idx):
                        field = FieldInfo(
                            field_id=f"para_{para_idx + 1}",
                            field_name=label,
                            field_type="paragraph",
                            location={
                                "paragraph_idx": para_idx + 1,
                                "input_type": "next_line"
                            },
                            description=char_limit,
                            max_length=int(re.search(r'\d+', char_limit).group()) if re.search(r'\d+', char_limit) else None
                        )
                        fields.append(field)
                    # 文字数制限がない場合は、サブ項目が続く可能性があるので
                    # 親質問として記憶するだけで、フィールドは作成しない
                    continue
                
                # 番号なしの通常テキストが来たら親質問コンテキストをリセット
                # ただし、サブ項目（①②等）が来る可能性があるので少し猶予を持たせる
                if parent_question and para_idx > parent_question_idx + 5:
                    parent_question = None
                    parent_question_idx = None
                
                # パターン3: 下線プレースホルダー
                underline_match = re.search(r'(.+?)[:：]\s*[_＿]{3,}', text)
                if underline_match:
                    field = FieldInfo(
                        field_id=f"para_{para_idx}",
                        field_name=underline_match.group(1).strip(),
                        field_type="paragraph",
                        location={
                            "paragraph_idx": para_idx,
                            "input_type": "underline"
                        }
                    )
                    fields.append(field)
                    continue
                
                # パターン4: 空括弧プレースホルダー
                bracket_match = re.search(r'(.+?)[(（]\s*[　\s]*[)）]', text)
                if bracket_match:
                    field = FieldInfo(
                        field_id=f"para_{para_idx}",
                        field_name=bracket_match.group(1).strip(),
                        field_type="paragraph",
                        location={
                            "paragraph_idx": para_idx,
                            "input_type": "bracket"
                        }
                    )
                    fields.append(field)
                    continue
                
                # パターン5: コロン終端（次行入力）
                colon_end_match = re.match(r'^(.+?)[:：]\s*$', text)
                if colon_end_match:
                    # 次がテーブルならスキップ
                    if is_next_element_table(para_idx):
                        continue

                    field = FieldInfo(
                        field_id=f"para_{para_idx + 1}",  # 次の段落
                        field_name=colon_end_match.group(1).strip(),
                        field_type="paragraph",
                        location={
                            "paragraph_idx": para_idx + 1,
                            "input_type": "next_line"
                        }
                    )
                    fields.append(field)
                    continue

                # パターン6: 見出し + 空行（記述式の可能性大）
                # 条件: 短いテキストで終わり、次の段落が空、かつキーワードを含む
                next_para_text = paragraphs[para_idx + 1].text.strip() if para_idx + 1 < len(paragraphs) else ""
                
                # キーワード判定（質問や項目名っぽいもの）
                is_question_header = any(k in text for k in ['について', '教えて', '概要', '内容', '理由', '目的', '計画', '戦略', '課題', '成果', '詳細'])
                
                # 末尾が「：」や「?」で終わる、または番号付き
                is_header_style = text.endswith(('：', ':', '？', '?')) or re.match(r'^[\d①-⑩]+[\.．]', text)
                
                if (is_question_header or is_header_style) and not next_para_text:
                    # 次がテーブルならスキップ
                    if is_next_element_table(para_idx):
                        continue

                    # 次の段落が空 = 入力スペースとみなす
                    field = FieldInfo(
                        field_id=f"para_{para_idx + 1}",
                        field_name=f"{text}（記述）",
                        field_type="paragraph",
                        location={
                            "paragraph_idx": para_idx + 1,
                            "input_type": "next_line_narrative"
                        },
                        description="200-600文字程度の詳細な記述",
                        input_length_type="long"
                    )
                    fields.append(field)
                    continue
        
        except Exception as e:
            self.logger.warning(f"[FORMAT_MAPPER] Error analyzing Word paragraphs: {e}")
        
        return fields
    
    def _enhance_fields_with_vlm(
        self, 
        fields: List[FieldInfo], 
        file_path: str, 
        file_type: str
    ) -> List[FieldInfo]:
        """VLMでフィールドの意味を推論して強化（チャンク分割対応）"""
        
        if not self.client or not fields:
            return fields
        
        try:
            # チャンクサイズの設定（1チャンクあたり最大50フィールド）
            CHUNK_SIZE = 50
            
            # フィールドをチャンクに分割
            field_chunks = [fields[i:i + CHUNK_SIZE] for i in range(0, len(fields), CHUNK_SIZE)]
            
            self.logger.info(f"[FORMAT_MAPPER] Enhancing {len(fields)} fields in {len(field_chunks)} chunks")
            
            all_enhanced_data = {}
            
            for chunk_idx, chunk in enumerate(field_chunks):
                try:
                    # フィールド情報をテキスト化
                    fields_text = "\n".join([
                        f"- {f.field_name} (type: {f.field_type}, id: {f.field_id})"
                        for f in chunk
                    ])
                    
                    prompt = f"""
以下は{file_type.upper()}申請フォーマットから検出されたフィールド一覧（チャンク {chunk_idx + 1}/{len(field_chunks)}）です。
各フィールドが実際に何を入力すべきかを推論し、JSONで返してください。

検出されたフィールド:
{fields_text}

出力形式（JSON配列）:
[
  {{
    "field_id": "フィールドID",
    "description": "このフィールドに入力すべき内容の説明",
    "required": true/false,
    "max_length": 数値またはnull,
    "input_length_type": "short" または "long"
  }},
  ...
]

■ input_length_type (入力タイプ) の判定基準:
- "short" (デフォルト): 
  - 基本的に全てのテーブル入力は "short" と判定してください。
  - 名前、日付、金額、電話番号、選択肢、短い文章など、1行で収まる内容。
- "long": 
  - 項目名に「概要」「説明」「理由」「目的」「計画」「内容」など、明らかに長文記述を求めている単語が含まれる場合のみ "long" と判定してください。

■ 重要: 数値項目（金額・数量など）の制約:
- 項目名が「金額」「価格」「数量」「人数」「単価」などの数値を示す場合:
- description には必ず「数値、空、または"未定"のみ入力可能」という旨を記載してください。
- 単位（円、人など）が分かる場合はそれも記載してください。

JSONのみを出力してください。
"""
                    
                    response = self.client.models.generate_content(
                        model=self.model_name,
                        contents=prompt
                    )
                    
                    # JSONをパース
                    response_text = response.text.strip()
                    # コードブロックを除去
                    if "```json" in response_text:
                        response_text = response_text.split("```json")[1].split("```")[0]
                    elif "```" in response_text:
                        response_text = response_text.split("```")[1].split("```")[0]
                    
                    enhanced_data = json.loads(response_text)
                    
                    # 結果をマージ
                    for item in enhanced_data:
                        all_enhanced_data[item["field_id"]] = item
                    
                    self.logger.info(f"[FORMAT_MAPPER] Chunk {chunk_idx + 1}/{len(field_chunks)}: enhanced {len(enhanced_data)} fields")
                    
                except Exception as chunk_error:
                    self.logger.warning(f"[FORMAT_MAPPER] Chunk {chunk_idx + 1} enhancement failed: {chunk_error}")
            
            # フィールドを更新
            for field in fields:
                if field.field_id in all_enhanced_data:
                    data = all_enhanced_data[field.field_id]
                    field.description = data.get("description")
                    field.required = data.get("required", False)
                    if data.get("max_length"):
                        field.max_length = data["max_length"]
                    # input_length_typeを反映（デフォルトはunknown）
                    input_length = data.get("input_length_type", "unknown")
                    if input_length in ["short", "long"]:
                        field.input_length_type = input_length
                    else:
                        # キーワードベースのフォールバック判定
                        field.input_length_type = self._infer_input_length_type(field.field_name)
            
            self.logger.info(f"[FORMAT_MAPPER] Enhanced total {len(all_enhanced_data)} fields with VLM")
            
        except Exception as e:
            self.logger.warning(f"[FORMAT_MAPPER] VLM enhancement failed: {e}")
        
        return fields
    
    def _infer_input_length_type(self, field_name: str) -> str:
        """
        フィールド名からinput_length_type（short/long）を推論する。
        VLM判定がない場合のフォールバック。
        """
        if not field_name:
            return "unknown"
        
        name_lower = field_name.lower()
        
        # 長文（long）を示すキーワード
        long_keywords = [
            '概要', '説明', '理由', '目的', '計画', '内容', '詳細',
            'プロジェクト', '事業', '活動', '取り組み', '実績', '成果',
            '課題', '問題', '背景', 'ビジョン', 'ミッション',
            'コメント', '備考', '補足', '自由記述',
            'description', 'summary', 'overview', 'detail', 'comment'
        ]
        
        # 短文（short）を示すキーワード
        short_keywords = [
            '名', '日', '年', '月', '番号', '電話', 'tel', 'fax',
            'url', 'メール', 'email', '住所', '金額', '円', '人数',
            '数', 'no', 'id', 'コード', '区分', '種別', '選択',
            'チェック', 'check'
        ]
        
        # 長文キーワードをチェック
        for kw in long_keywords:
            if kw in name_lower:
                return "long"
        
        # 短文キーワードをチェック
        for kw in short_keywords:
            if kw in name_lower:
                return "short"
        
        return "unknown"
    
    def _calculate_field_importance(self, field: FieldInfo, field_index: int = 0) -> int:
        """
        フィールドの重要度をスコアリングする。
        
        Args:
            field: 評価対象のフィールド
            field_index: フィールドの順序（0から始まる）
            
        Returns:
            重要度スコア（高いほど重要）
        """
        score = 0
        field_name_lower = field.field_name.lower()
        
        # 1. 必須フラグ: +50点
        if field.required:
            score += 50
        
        # 2. キーワードマッチング
        # 最優先キーワード: +30点
        high_priority_keywords = ['団体名', '法人名', '代表者', '連絡先', '電話', 'tel', 'email', 'メールアドレス', '住所']
        for kw in high_priority_keywords:
            if kw in field_name_lower:
                score += 30
                break
        
        # 重要キーワード: +20点
        important_keywords = ['事業名', '目的', '概要', 'プロジェクト', '計画', '活動', '取り組み']
        for kw in important_keywords:
            if kw in field_name_lower:
                score += 20
                break
        
        # 金額関連: +15点
        amount_keywords = ['金額', '予算', '費用', 'cost', '円']
        for kw in amount_keywords:
            if kw in field_name_lower:
                score += 15
                break
        
        # 3. フィールドタイプ
        # 長文フィールド（詳細な記述が必要）: +10点
        if field.input_length_type == "long":
            score += 10
        # 短文フィールド: +5点
        elif field.input_length_type == "short":
            score += 5
        
        # 4. 位置による加点（最初の方のフィールド）
        # 先頭から30フィールド以内: 位置に応じて最大+5点
        if field_index < 30:
            score += max(5 - (field_index // 6), 0)
        
        # 5. 文字数制限がある場合: +3点（正式な制限があるフィールドは重要度が高い）
        if field.max_length and field.max_length > 0:
            score += 3
        
        return score
    
    def _limit_fields_by_importance(self, fields: List[FieldInfo], max_fields: int = 50) -> Tuple[List[FieldInfo], int]:
        """
        フィールドを重要度順にソートし、上限数に制限する。
        
        Args:
            fields: フィールドのリスト
            max_fields: 最大フィールド数（デフォルト: 50）
            
        Returns:
            (制限後のフィールドリスト, 元のフィールド数)
        """
        original_count = len(fields)
        
        # フィールド数が上限以下の場合はそのまま返す
        if original_count <= max_fields:
            self.logger.info(f"[FORMAT_MAPPER] Field count ({original_count}) is within limit ({max_fields})")
            return fields, original_count
        
        # 重要度スコアを計算
        field_scores = []
        for idx, field in enumerate(fields):
            score = self._calculate_field_importance(field, idx)
            field_scores.append((field, score))
            self.logger.debug(f"[FORMAT_MAPPER] Field '{field.field_name}' (ID: {field.field_id}): importance score = {score}")
        
        # スコアでソート（降順）
        field_scores.sort(key=lambda x: x[1], reverse=True)
        
        # 上位max_fieldsフィールドを選択
        selected_fields = [field for field, score in field_scores[:max_fields]]
        
        # ログ出力
        self.logger.warning(
            f"[FORMAT_MAPPER] Field limit applied: {original_count} fields found, "
            f"limited to top {max_fields} by importance. "
            f"Skipped {original_count - max_fields} fields."
        )
        
        # スキップされたフィールドの例を出力（デバッグ用）
        skipped_fields = [field for field, score in field_scores[max_fields:max_fields+5]]
        if skipped_fields:
            skipped_names = [f.field_name for f in skipped_fields]
            self.logger.info(f"[FORMAT_MAPPER] Examples of skipped fields: {', '.join(skipped_names[:5])}")
        
        return selected_fields, original_count

    
    def map_draft_to_fields(
        self, 
        fields: List[FieldInfo], 
        draft: str,
        grant_name: str = "",
        include_field_info: bool = True
    ) -> Dict[str, Any]:
        """
        Geminiでドラフト内容を各フィールドにマッピングする。
        
        Args:
            fields: フィールド情報のリスト
            draft: 生成されたドラフトテキスト
            grant_name: 助成金名
            include_field_info: フィールド情報（入力パターン等）を含むかどうか
            
        Returns:
            Dict[str, Any] - 入力値と入力パターン情報を含むマッピング
            例: {
                "para_5": {
                    "value": "入力値",
                    "input_pattern": "inline",
                    "field_name": "団体名"
                }
            }
            include_field_info=Falseの場合は従来通り {field_id: 入力値}
        """
        if not fields or not draft or not self.client:
            return {}
        
        try:
            # フィールド数制限（大量の項目がある場合は重要度順に選択）
            limited_fields, original_count = self._limit_fields_by_importance(fields, max_fields=50)
            skipped_count = original_count - len(limited_fields)
            
            # 制限情報を保存（drafter側で参照可能にする）
            self.last_total_field_count = original_count
            self.last_skipped_field_count = skipped_count
            
            # フィールド情報をプロンプト用にフォーマット
            fields_prompt = "\n".join([
                f"- field_id: {f.field_id}\n  名前: {f.field_name}\n  説明: {f.description or '不明'}\n  文字数制限: {f.max_length or '制限なし'}\n  入力タイプ: {f.input_length_type}"
                for f in limited_fields
            ])
            
            prompt = f"""
以下のドラフト内容を、申請フォーマットの各フィールドに適切にマッピングしてください。

## 対象助成金
{grant_name}

## ドラフト内容
{draft[:10000]}

## マッピング先フィールド
{fields_prompt}

## 出力形式（JSON）
{{
  "field_id_1": "このフィールドに入力する内容",
  "field_id_2": "このフィールドに入力する内容",
  ...
}}

注意:
- 各フィールドの文字数制限を守ってください
- ドラフトに該当する内容がない場合は空文字列にしてください
- JSONのみを出力してください

■ 重要: 計算処理
- 「小計」や「合計」という名前のフィールドは、あなたが入力した他の数値フィールド（単価×数量など）の合計値を**必ず計算してその値を記入**してください。空欄にしてはいけません。

■ 重要: 長文記述
- 入力タイプが "long" のフィールド（「概要」「理由」「計画」など）は、ドラフトの内容を元に、必ず**200〜600文字程度の具体的で詳細な文章**を生成してください。
- "TBD" や "詳細は後日決定" などの短文で逃げないでください。不足している情報は、それらしい内容を補完してでもドラフトを埋めてください。

■ 重要: 短文記述（テーブル内など）
- 入力タイプが "short" のフィールドについては、決して「...」や「省略」を使わないでください。
- 長くなる場合は、**短い単語、体言止め、またはコマンド的な表現**（例：「Xを実施」「Yを確認」）に言い換えて、必ず枠内に収まる表現を生成してください。
"""
            
            response = self.client.models.generate_content(
                model=self.model_name,
                contents=prompt
            )
            
            # JSONをパース
            response_text = response.text.strip()
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0]
            elif "```" in response_text:
                response_text = response_text.split("```")[1].split("```")[0]
            
            mapping = json.loads(response_text)
            
            self.logger.info(f"[FORMAT_MAPPER] Mapped {len(mapping)} fields from draft")
            
            # 入力パターン情報を含める場合
            if include_field_info:
                # フィールドIDをキーにしたlocation/pattern情報を作成
                field_info_map = {f.field_id: f for f in fields}
                
                enhanced_mapping = {}
                for field_id, value in mapping.items():
                    field_info = field_info_map.get(field_id)
                    if field_info:
                        enhanced_mapping[field_id] = {
                            "value": value,
                            "input_pattern": field_info.location.get("input_pattern", "inline"),
                            "field_name": field_info.field_name,
                            "field_type": field_info.field_type,
                            "location": field_info.location
                        }
                    else:
                        enhanced_mapping[field_id] = {
                            "value": value,
                            "input_pattern": "inline",
                            "field_name": "",
                            "field_type": "unknown",
                            "location": {}
                        }
                
                return enhanced_mapping
            
            return mapping
            
        except Exception as e:
            self.logger.error(f"[FORMAT_MAPPER] Draft mapping failed: {e}")
            return {}
    
    def fill_fields_individually(
        self,
        fields: List[FieldInfo],
        profile: str,
        grant_name: str = "",
        grant_info: str = "",
        progress_callback=None,
        refined_draft: str = ""
    ) -> Dict[str, Any]:
        """
        各フィールドに対して個別にGemini 3.0 Flashを呼び出し、
        プロファイル情報をもとに回答を生成する。
        
        Args:
            fields: フィールド情報のリスト
            profile: Soul Profile（NPOのプロファイル情報）
            grant_name: 助成金名
            grant_info: 助成金の詳細情報（募集要項など）
            progress_callback: 進捗通知用コールバック関数 (field_idx, total, field_name) -> None
            refined_draft: 推敲ループで改善されたドラフト（参照情報として使用）
            
        Returns:
            Dict[str, Any] - 入力値と入力パターン情報を含むマッピング
            例: {
                "para_5": {
                    "value": "入力値",
                    "input_pattern": "inline",
                    "field_name": "団体名"
                }
            }
        """
        if not fields or not profile or not self.client:
            self.logger.warning("[FORMAT_MAPPER] Missing required inputs for fill_fields_individually")
            return {}
        
        # gemini-3.0-flashモデルを使用（項目別処理用）
        flash_model = "gemini-3-flash-preview"
        
        result = {}
        total_fields = len(fields)
        
        # 推敲済みドラフトの参照セクションを構築
        draft_reference = ""
        if refined_draft:
            draft_reference = f"""
# 推敲済みドラフト（参照用）
以下は既に推敲ループで改善されたドラフトです。この内容を参考に回答を作成してください。
ドラフトと矛盾しない内容で、かつドラフトの表現スタイルを踏襲してください。

{refined_draft[:5000]}

"""
        
        for idx, field in enumerate(fields):
            try:
                # 進捗通知
                if progress_callback:
                    progress_callback(idx, total_fields, field.field_name)
                
                self.logger.info(f"[FORMAT_MAPPER] Processing field {idx + 1}/{total_fields}: {field.field_name}")
                
                # 入力長さタイプに応じた指示を構築
                length_instruction = ""
                if field.input_length_type == "short":
                    length_instruction = "\n\n**回答形式**: 1行以内の簡潔な回答（名前、日付、数値など）"
                    if field.max_length:
                        length_instruction += f"\n**重要**: {field.max_length}文字以内で**完全に完結する**文章やフレーズにしてください。\n- **禁止事項**: 文末を「...」や「…」で省略することは**厳禁**です。\n- 文末が切れないように体言止めなどを活用し、必ず制限内に収めてください。"
                elif field.input_length_type == "long":
                    if field.max_length:
                        # 長文で文字数制限あり：制限値の90%を目標に
                        target_length = int(field.max_length * 0.9)
                        length_instruction = f"""
**文字数制限**: {field.max_length}字以内
**目標文字数**: {target_length}字程度（制限の90%を目安に）
**重要**: 文字数制限を厳守してください。{field.max_length}字を超えないよう、適切に要約してください。
**禁止事項**: 文末を「...」で省略してはいけません。文章を短く書き直して収めてください。"""
                    else:
                        length_instruction = "\n\n**回答形式**: 詳細な長文回答（100〜500字程度）"
                else:
                    # unknownの場合
                    if field.max_length:
                        length_instruction = f"\n\n**文字数制限**: {field.max_length}字以内で回答してください。文章が途中で切れないようにしてください。"
                
                prompt = f"""あなたはNPOの助成金申請書作成を支援する専門家です。
以下のNPOプロファイル情報と助成金情報に基づいて、申請書の指定された項目に対する回答を作成してください。

# NPO Soul Profile（魂のプロファイル）
{profile[:6000]}

# 対象助成金
助成金名: {grant_name}
{grant_info[:3000] if grant_info else ""}
{draft_reference}
# 回答すべき項目
- **項目名**: {field.field_name}
- **説明**: {field.description or "（説明なし）"}
- **入力形式**: {field.input_length_type}
{length_instruction}

# 指示
1. NPOの情報を最大限活用して、この項目に対する適切な回答を作成してください
2. 助成金の目的や評価基準を意識した回答を心がけてください
3. 具体的で説得力のある内容にしてください
4. 回答本文のみを出力してください（項目名や説明は不要）
5. 文字数制限がある場合は必ず守ってください。**制限超過による途中切れや「...」での省略は厳禁**です。

# 懸念点の報告（重要）
- プロファイルに情報がなく回答できない場合: 回答の末尾に `[MISSING_INFO: 理由]` を付記
- 回答に自信がない・確認が必要な場合: 回答の末尾に `[UNCERTAIN: 理由]` を付記
- 問題なく回答できる場合: タグは不要

# 回答:"""

                response = self.client.models.generate_content(
                    model=flash_model,
                    contents=prompt
                )
                
                value = response.text.strip()
                
                # 懸念点タグを抽出
                concern_type = "none"
                concern_reason = ""
                
                import re
                # [MISSING_INFO: 理由] パターンを検出
                missing_match = re.search(r'\[MISSING_INFO:\s*([^\]]+)\]', value)
                if missing_match:
                    concern_type = "missing_info"
                    concern_reason = missing_match.group(1).strip()
                    value = re.sub(r'\s*\[MISSING_INFO:[^\]]+\]', '', value).strip()
                    self.logger.info(f"[FORMAT_MAPPER] Field {field.field_name} has missing info: {concern_reason}")
                
                # [UNCERTAIN: 理由] パターンを検出
                uncertain_match = re.search(r'\[UNCERTAIN:\s*([^\]]+)\]', value)
                if uncertain_match:
                    concern_type = "uncertain"
                    concern_reason = uncertain_match.group(1).strip()
                    value = re.sub(r'\s*\[UNCERTAIN:[^\]]+\]', '', value).strip()
                    self.logger.info(f"[FORMAT_MAPPER] Field {field.field_name} is uncertain: {concern_reason}")
                
                # 文字数制限チェック & 自動リトライ
                max_retries = 2
                retry_count = 0
                original_value = value
                
                while field.max_length and len(value) > field.max_length and retry_count < max_retries:
                    retry_count += 1
                    overflow_chars = len(value) - field.max_length
                    self.logger.warning(
                        f"[FORMAT_MAPPER] Field {field.field_id} exceeded max length "
                        f"({len(value)} > {field.max_length}, over by {overflow_chars}). Retry {retry_count}/{max_retries}"
                    )
                    
                    # 短縮リトライプロンプト
                    retry_prompt = f"""前回の回答が{field.max_length}字の制限を{overflow_chars}字オーバーしました。
以下の回答を**必ず{field.max_length}字以内**に短縮してください。

# 短縮対象の回答
{value}

# 短縮の指示
1. 制限の90%である{int(field.max_length * 0.9)}字程度を目標にしてください
2. 重要な情報を残しつつ、冗長な表現を削除してください
3. 体言止めや簡潔な表現を活用してください
4. 「...」「…」での省略は禁止です
5. 短縮後の文章のみを出力してください

# 短縮後の回答:"""

                    try:
                        retry_response = self.client.models.generate_content(
                            model=flash_model,
                            contents=retry_prompt
                        )
                        value = retry_response.text.strip()
                        self.logger.info(
                            f"[FORMAT_MAPPER] Retry {retry_count}: shortened to {len(value)} chars "
                            f"(limit: {field.max_length})"
                        )
                    except Exception as retry_error:
                        self.logger.error(f"[FORMAT_MAPPER] Retry failed: {retry_error}")
                        break
                
                # 最終的な文字数超過チェック
                if field.max_length and len(value) > field.max_length:
                    overflow_chars = len(value) - field.max_length
                    overflow_percentage = (overflow_chars / field.max_length) * 100
                    self.logger.warning(
                        f"[FORMAT_MAPPER] Field {field.field_id} still exceeds limit after retries "
                        f"({len(value)} > {field.max_length}, {overflow_percentage:.1f}% over)"
                    )
                    # 懸念点として記録（既存の懸念がなければ）
                    if concern_type == "none":
                        concern_type = "length_exceeded"
                        concern_reason = f"{field.max_length}字制限を{overflow_chars}字超過（{overflow_percentage:.1f}%オーバー）"
                
                # 省略記号検知
                if value.endswith("...") or value.endswith("…"):
                    self.logger.warning(f"[FORMAT_MAPPER] Field {field.field_id} ends with ellipsis")
                    if concern_type == "none":
                        concern_type = "truncated"
                        concern_reason = "回答が省略記号で終わっています。手動で修正が必要です。"
                
                # 結果を格納（懸念点情報を含む）
                result[field.field_id] = {
                    "value": value,
                    "input_pattern": field.location.get("input_pattern", "inline"),
                    "field_name": field.field_name,
                    "field_type": field.field_type,
                    "location": field.location,
                    "input_length_type": field.input_length_type,
                    "concern_type": concern_type,
                    "concern_reason": concern_reason,
                    "max_length": field.max_length,
                    "actual_length": len(value),
                    "retry_count": retry_count
                }
                
                self.logger.info(f"[FORMAT_MAPPER] Successfully filled field: {field.field_name} ({len(value)} chars, concern: {concern_type})")
                
            except Exception as e:
                self.logger.error(f"[FORMAT_MAPPER] Error filling field {field.field_name}: {e}")
                # エラーの場合は空の値を設定
                result[field.field_id] = {
                    "value": "",
                    "input_pattern": field.location.get("input_pattern", "inline"),
                    "field_name": field.field_name,
                    "field_type": field.field_type,
                    "location": field.location,
                    "error": str(e)
                }
        
        self.logger.info(f"[FORMAT_MAPPER] Completed filling {len(result)} fields individually")
        return result
    
    def analyze_format_file(self, file_path: str) -> Tuple[List[FieldInfo], str]:
        """
        ファイル形式を自動判定してフィールドを解析する。
        
        Args:
            file_path: ファイルパス
            
        Returns:
            (フィールドリスト, ファイルタイプ)
        """
        import os
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext in ['.xlsx', '.xlsm', '.xls']:
            return self.analyze_excel_fields(file_path), "excel"
        elif ext in ['.docx', '.doc']:
            return self.analyze_word_fields(file_path), "word"
        else:
            self.logger.warning(f"[FORMAT_MAPPER] Unsupported file type: {ext}")
            return [], "unknown"
    
    def generate_concern_report(self, field_values: Dict[str, Any]) -> str:
        """
        入力結果から懸念点レポートを生成する（強化版）。
        
        Args:
            field_values: fill_fields_individuallyの戻り値
            
        Returns:
            Markdown形式のレポート（品質スコア付き）
        """
        missing_info_fields = []
        uncertain_fields = []
        length_exceeded_fields = []
        truncated_fields = []
        
        # 品質スコア計算用
        total_fields = len(field_values)
        good_fields = 0
        retry_total = 0
        
        for field_id, data in field_values.items():
            concern_type = data.get("concern_type", "none")
            concern_reason = data.get("concern_reason", "")
            field_name = data.get("field_name", field_id)
            retry_count = data.get("retry_count", 0)
            max_length = data.get("max_length")
            actual_length = data.get("actual_length", 0)
            
            retry_total += retry_count
            
            if concern_type == "none":
                good_fields += 1
            elif concern_type == "missing_info":
                missing_info_fields.append({
                    "field_name": field_name,
                    "reason": concern_reason
                })
            elif concern_type == "uncertain":
                uncertain_fields.append({
                    "field_name": field_name,
                    "reason": concern_reason
                })
            elif concern_type == "length_exceeded":
                length_exceeded_fields.append({
                    "field_name": field_name,
                    "reason": concern_reason,
                    "max_length": max_length,
                    "actual_length": actual_length
                })
            elif concern_type == "truncated":
                truncated_fields.append({
                    "field_name": field_name,
                    "reason": concern_reason
                })
        
        # 品質スコアを計算（0-100）
        if total_fields > 0:
            quality_score = int((good_fields / total_fields) * 100)
            completion_rate = int(((total_fields - len(missing_info_fields)) / total_fields) * 100)
        else:
            quality_score = 0
            completion_rate = 0
        
        # 懸念点があるか確認
        concerns_exist = (missing_info_fields or uncertain_fields or 
                         length_exceeded_fields or truncated_fields)
        concern_total = len(missing_info_fields) + len(uncertain_fields) + len(length_exceeded_fields) + len(truncated_fields)
        
        # 簡潔な統計サマリー形式のレポート
        report = "## 📋 ドラフト品質サマリー\n\n"
        
        # 品質スコアの表示
        if quality_score >= 80:
            score_emoji = "🟢"
            score_label = "優良"
        elif quality_score >= 60:
            score_emoji = "🟡"
            score_label = "標準"
        else:
            score_emoji = "🔴"
            score_label = "要改善"
        
        report += f"- **品質スコア**: {score_emoji} **{quality_score}点** ({score_label})\n"
        report += f"- ✅ **正常入力**: {good_fields}項目\n"
        
        # 懸念点がない場合
        if not concerns_exist:
            if retry_total > 0:
                report += f"- 🔧 **自動修正**: {retry_total}回の文字数超過を修正\n"
            report += "\n> すべての項目が正常に入力されました！\n"
            return report
        
        # 懸念点がある場合の統計サマリー
        report += f"- ⚠️ **要確認**: {concern_total}項目（Word/Excelのコメントを参照）\n"
        
        # 内訳を簡潔に表示
        breakdown = []
        if missing_info_fields:
            breakdown.append(f"情報不足: {len(missing_info_fields)}項目")
        if uncertain_fields:
            breakdown.append(f"不確実: {len(uncertain_fields)}項目")
        if length_exceeded_fields:
            breakdown.append(f"文字数超過: {len(length_exceeded_fields)}項目")
        if truncated_fields:
            breakdown.append(f"省略あり: {len(truncated_fields)}項目")
        
        if breakdown:
            report += f"  - {' / '.join(breakdown)}\n"
        
        if retry_total > 0:
            report += f"- 🔧 **自動修正**: {retry_total}回の文字数超過を修正\n"
        
        report += "\n> 💡 詳細は添付のWord/Excelファイル内のコメントをご確認ください。\n"
        
        return report


