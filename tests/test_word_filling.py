"""
Test script for Word document filling functionality.
Tests the format_field_mapper and document_filler modules.
"""

import sys
import os

# Add project root to path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import asyncio
import logging

logging.basicConfig(level=logging.DEBUG)

def test_word_filling(test_file_path: str):
    """Test Word document field detection and filling."""
    
    print(f"\n{'='*60}")
    print(f"Testing Word document filling")
    print(f"File: {test_file_path}")
    print(f"{'='*60}\n")
    
    # Check if file exists
    if not os.path.exists(test_file_path):
        print(f"ERROR: File not found: {test_file_path}")
        return False
    
    try:
        # Step 1: Test FormatFieldMapper
        print("Step 1: Testing FormatFieldMapper...")
        from src.logic.format_field_mapper import FormatFieldMapper
        
        mapper = FormatFieldMapper()
        fields, file_type = mapper.analyze_format_file(test_file_path)
        
        print(f"  File type detected: {file_type}")
        print(f"  Fields found: {len(fields)}")
        
        if fields:
            for i, field in enumerate(fields[:10]):  # Show first 10
                print(f"    [{i}] ID: {field.field_id}")
                print(f"        Name: {field.field_name}")
                print(f"        Type: {field.field_type}")
                print(f"        Location: {field.location}")
                print()
        else:
            print("  WARNING: No fields detected!")
            print("  This means the Word document structure isn't being recognized.")
            print()
            
            # Try to show document structure for debugging
            print("  Debugging document structure:")
            try:
                from docx import Document
                doc = Document(test_file_path)
                
                print(f"    Paragraphs: {len(doc.paragraphs)}")
                for i, para in enumerate(doc.paragraphs[:5]):
                    print(f"      [{i}] {para.text[:80] if para.text else '(empty)'}...")
                
                print(f"    Tables: {len(doc.tables)}")
                for i, table in enumerate(doc.tables[:3]):
                    print(f"      Table {i}: {len(table.rows)} rows x {len(table.rows[0].cells) if table.rows else 0} cols")
                    for row_idx, row in enumerate(table.rows[:3]):
                        for col_idx, cell in enumerate(row.cells):
                            text = cell.text.strip()[:30] if cell.text else "(empty)"
                            print(f"        [{row_idx},{col_idx}]: {text}")
            except Exception as e:
                print(f"    Error reading document: {e}")
        
        # Step 2: Test DocumentFiller (even without detected fields)
        print("\nStep 2: Testing DocumentFiller with sample data...")
        from src.tools.document_filler import DocumentFiller
        
        filler = DocumentFiller(output_dir="./testdata/output")
        
        # Create sample field values based on typical application form
        sample_values = {
            "table0_0_1": "テスト団体名",
            "table0_1_1": "テスト代表者名",
            "table0_2_1": "東京都渋谷区テスト1-2-3",
            "para_0": "テスト事業の概要説明文です。",
        }
        
        print(f"  Sample values: {sample_values}")
        
        output_path, message = filler.fill_document(
            test_file_path,
            sample_values,
            user_id="test_user"
        )
        
        if output_path:
            print(f"\n  SUCCESS!")
            print(f"  Output file: {output_path}")
            print(f"  Message: {message}")
            
            # Verify output file exists
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                print(f"  File size: {file_size} bytes")
                return True
            else:
                print(f"  WARNING: Output file was not created")
                return False
        else:
            print(f"\n  FAILED: {message}")
            return False
            
    except ImportError as e:
        print(f"\nIMPORT ERROR: {e}")
        print("Make sure you have installed: pip install python-docx openpyxl")
        return False
    except Exception as e:
        print(f"\nERROR: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == "__main__":
    # Default test file path
    test_file = sys.argv[1] if len(sys.argv) > 1 else "testdata/application.docx"
    
    success = test_word_filling(test_file)
    
    print(f"\n{'='*60}")
    print(f"TEST RESULT: {'PASSED' if success else 'FAILED'}")
    print(f"{'='*60}")
    
    sys.exit(0 if success else 1)
